import asyncio
import base64
import ipaddress
import re
import socket
import zipfile
import xml.etree.ElementTree as ET
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from urllib.parse import ParseResult, urljoin, urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageOps
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFError, TTFont
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.platypus import Image as PDFImage


def extract_document_text(filename: str, data: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1]
    if suffix == "pdf":
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if suffix in {"docx", "doc"}:
        if suffix == "doc":
            raise ValueError("暂不支持旧版 .doc，请另存为 .docx 后上传")
        document = Document(BytesIO(data))
        # Many formal resume templates put all visible text in tables.  The
        # old implementation read only top-level paragraphs and treated those
        # files as blank.
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(paragraph.text for paragraph in cell.paragraphs)
        extracted = "\n".join(part.strip() for part in parts if part and part.strip()).strip()
        return extracted or _extract_docx_xml_text(data)
    if suffix in {"txt", "md"}:
        return data.decode("utf-8", errors="replace").strip()
    raise ValueError("仅支持 PDF、DOCX、TXT 文件")


def _extract_docx_xml_text(data: bytes) -> str:
    """Read DOCX textboxes, nested tables and headers missed by python-docx."""
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    lines: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            names = [name for name in archive.namelist() if name == "word/document.xml" or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)]
            for name in names:
                root = ET.fromstring(archive.read(name))
                for paragraph in root.iter(f"{namespace}p"):
                    values: list[str] = []
                    for node in paragraph.iter():
                        if node.tag == f"{namespace}t" and node.text:
                            values.append(node.text)
                        elif node.tag in {f"{namespace}tab", f"{namespace}br", f"{namespace}cr"}:
                            values.append(" ")
                    line = "".join(values).strip()
                    if line:
                        lines.append(line)
    except (ET.ParseError, OSError, zipfile.BadZipFile):
        return ""
    return "\n".join(dict.fromkeys(lines)).strip()


def _is_public_address(address: str) -> bool:
    try:
        ip = ipaddress.ip_address(address.split("%", 1)[0])
    except ValueError:
        return False
    return ip.is_global and not ip.is_multicast


async def _resolve_public_addresses(url: str) -> tuple[ParseResult, str, int, list[str]]:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("仅支持可公开访问的 HTTP/HTTPS 岗位链接")
    try:
        hostname = parsed.hostname.encode("idna").decode("ascii")
        port = parsed.port or (443 if parsed.scheme == "https" else 80)
        resolved = await asyncio.to_thread(socket.getaddrinfo, hostname, port, type=socket.SOCK_STREAM)
    except (socket.gaierror, UnicodeError, ValueError) as exc:
        raise ValueError("仅支持可公开访问的 HTTP/HTTPS 岗位链接") from exc
    addresses = list(dict.fromkeys(item[4][0].split("%", 1)[0] for item in resolved))
    if not addresses or any(not _is_public_address(address) for address in addresses):
        raise ValueError("仅支持可公开访问的 HTTP/HTTPS 岗位链接")
    return parsed, hostname, port, addresses


def _pinned_target(parsed: ParseResult, hostname: str, port: int, address: str) -> tuple[str, str, dict[str, str]]:
    ip = ipaddress.ip_address(address)
    ip_host = f"[{ip}]" if ip.version == 6 else str(ip)
    default_port = 443 if parsed.scheme == "https" else 80
    netloc = ip_host if port == default_port else f"{ip_host}:{port}"
    pinned_url = parsed._replace(netloc=netloc, fragment="").geturl()
    original_host = f"[{hostname}]" if ":" in hostname else hostname
    host_header = original_host if port == default_port else f"{original_host}:{port}"
    extensions = {"sni_hostname": hostname} if parsed.scheme == "https" else {}
    return pinned_url, host_header, extensions


def _response_for_original_url(response: httpx.Response, url: str, content: bytes = b"") -> httpx.Response:
    headers = response.headers.copy()
    for name in ("content-encoding", "content-length", "transfer-encoding"):
        headers.pop(name, None)
    return httpx.Response(
        status_code=response.status_code,
        headers=headers,
        content=content,
        request=httpx.Request("GET", url),
    )


async def _request_pinned_page(
    url: str, parsed: ParseResult, hostname: str, port: int, address: str
) -> httpx.Response:
    pinned_url, host_header, extensions = _pinned_target(parsed, hostname, port, address)
    headers = {"User-Agent": "Mozilla/5.0 (compatible; ResumeAI/1.0)", "Host": host_header}
    async with httpx.AsyncClient(timeout=15, follow_redirects=False, headers=headers, trust_env=False) as client:
        async with client.stream("GET", pinned_url, extensions=extensions) as response:
            if response.is_redirect or not response.is_success:
                return _response_for_original_url(response, url)
            content_type = response.headers.get("content-type", "").lower()
            if "html" not in content_type:
                raise ValueError("岗位链接没有返回网页内容")
            content_length = response.headers.get("content-length")
            if content_length and content_length.isdigit() and int(content_length) > 2_000_000:
                raise ValueError("岗位页面超过 2MB，无法安全解析")
            content = bytearray()
            async for chunk in response.aiter_bytes():
                if len(content) + len(chunk) > 2_000_000:
                    raise ValueError("岗位页面超过 2MB，无法安全解析")
                content.extend(chunk)
            return _response_for_original_url(response, url, bytes(content))


async def _fetch_public_page(url: str) -> httpx.Response:
    current_url = url
    for _ in range(4):
        parsed, hostname, port, addresses = await _resolve_public_addresses(current_url)
        response = None
        last_request_error: httpx.RequestError | None = None
        timeout_error: httpx.TimeoutException | None = None
        for address in addresses:
            try:
                response = await _request_pinned_page(current_url, parsed, hostname, port, address)
                break
            except httpx.TimeoutException as exc:
                timeout_error = exc
                last_request_error = exc
            except httpx.RequestError as exc:
                last_request_error = exc
        if response is None:
            if timeout_error:
                raise ValueError("读取岗位页面超时，请稍后重试或改用文本粘贴、截图上传") from timeout_error
            if last_request_error:
                raise last_request_error
            raise ValueError("未能连接岗位页面")
        if response.is_redirect:
            location = response.headers.get("location")
            if not location:
                raise ValueError("岗位链接返回了无效跳转")
            current_url = urljoin(current_url, location)
            continue
        response.raise_for_status()
        return response
    raise ValueError("岗位链接跳转次数过多")


def _numbered_items(value: str) -> list[str]:
    cleaned = value.strip().strip("；;")
    if not cleaned:
        return []
    parts = re.split(r"(?:^|\n)\s*\d+\s*[.、]\s*", cleaned)
    return [part.strip().strip("；;") for part in parts if part.strip().strip("；;")]


JD_CHALLENGE_ERROR = "目标网站返回人机验证页，未获取到岗位正文。请改用岗位截图或复制岗位正文粘贴。"
JD_EMPTY_ERROR = "岗位解析结果为空，未识别到职位、公司、职责或要求。请改用岗位截图或复制岗位正文粘贴。"


def _text_items(value) -> list[str]:
    if isinstance(value, str):
        values = [value]
    elif isinstance(value, list):
        values = value
    else:
        values = []
    return [str(item).strip() for item in values if str(item or "").strip()]


def _looks_like_verification_page(value: str) -> bool:
    text = str(value or "")
    if not text:
        return False
    lower = text.lower()
    strong_markers = (
        "aliyun_waf",
        "cf_app_waf",
        "captcha",
        "geetest",
        "nc_token",
        "acw_sc__v2",
        "acw_tc",
    )
    chinese_markers = ("为了更好的访问体验，请进行验证", "请进行验证", "人机验证", "安全验证", "滑块验证")
    return any(marker in lower for marker in strong_markers) or any(marker in text for marker in chinese_markers)


def validate_jd_result(result: dict | None) -> dict:
    if not isinstance(result, dict):
        raise ValueError(JD_EMPTY_ERROR)
    raw_text = str(result.get("raw_text") or "")
    if _looks_like_verification_page(raw_text):
        raise ValueError(JD_CHALLENGE_ERROR)
    title = str(result.get("title") or "").strip()
    company = str(result.get("company") or "").strip()
    responsibilities = _text_items(result.get("responsibilities"))
    requirements = _text_items(result.get("requirements"))
    preferred = _text_items(result.get("preferred"))
    keywords = _text_items(result.get("keywords"))
    signal_count = len(responsibilities) + len(requirements) + len(preferred) + len(keywords)
    if not title and not company and signal_count < 2:
        raise ValueError(JD_EMPTY_ERROR)
    if not title and not (responsibilities or requirements):
        raise ValueError(JD_EMPTY_ERROR)
    return result


def _extract_gxrc_job(soup: BeautifulSoup, url: str) -> dict | None:
    hostname = (urlparse(url).hostname or "").lower()
    if hostname != "gxrc.com" and not hostname.endswith(".gxrc.com"):
        return None
    description = soup.select_one("pre#examineSensitiveWordsContent")
    if not description:
        return None
    raw_description = description.get_text("\n", strip=True)
    responsibilities_text = ""
    requirements_text = raw_description
    if "岗位职责：" in raw_description:
        requirements_text, responsibilities_text = raw_description.split("岗位职责：", 1)
    requirements_text = requirements_text.replace("任职要求：", "", 1).strip()
    title_node = soup.find("h1")
    company_node = soup.select_one(".ent-name a")
    title = title_node.get_text(" ", strip=True) if title_node else ""
    company = company_node.get_text(" ", strip=True) if company_node else ""
    if not title:
        page_title = soup.title.get_text(" ", strip=True) if soup.title else ""
        title = page_title.split("职位信息_", 1)[0].strip()
        if not company and "职位信息_" in page_title:
            company = page_title.split("职位信息_", 1)[1].split(" - ", 1)[0].strip()
    responsibilities = _numbered_items(responsibilities_text)
    requirements = _numbered_items(requirements_text)
    if not title or not (responsibilities or requirements):
        return None
    keywords = [node.get_text(" ", strip=True) for node in soup.select(".layui-keyword-text")]
    raw_text = "\n".join(part for part in (title, company, raw_description) if part)
    return {
        "title": title,
        "company": company,
        "responsibilities": responsibilities,
        "requirements": requirements,
        "preferred": [],
        "keywords": list(dict.fromkeys(keyword for keyword in keywords if keyword)),
        "raw_text": raw_text,
        "parser": "gxrc",
    }


async def fetch_job_page(url: str) -> tuple[str, dict | None]:
    parsed_url = urlparse(url)
    if (parsed_url.hostname or "").lower() == "m.gxrc.com":
        url = parsed_url._replace(netloc="www.gxrc.com").geturl()
    response = await _fetch_public_page(url)
    if _looks_like_verification_page(response.text):
        raise ValueError(JD_CHALLENGE_ERROR)
    soup = BeautifulSoup(response.text, "html.parser")
    structured = _extract_gxrc_job(soup, str(response.url))
    for element in soup(["script", "style", "nav", "footer", "noscript", "svg"]):
        element.decompose()
    text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
    if _looks_like_verification_page(text):
        raise ValueError(JD_CHALLENGE_ERROR)
    if len(text) < 30:
        raise ValueError("未能从链接提取到有效岗位内容，请改用文本粘贴或截图")
    return text[:50000], structured


async def fetch_job_text(url: str) -> str:
    text, _ = await fetch_job_page(url)
    return text


CONTACT_LABELS = {
    "age": "年龄",
    "birth_age": "年龄",
    "birth_date": "出生年月",
    "birthday": "出生日期",
    "birth_year": "出生年份",
    "gender": "性别",
    "phone": "电话",
    "mobile": "手机",
    "email": "邮箱",
    "location": "所在地",
    "city": "所在地",
    "address": "地址",
    "wechat": "微信",
    "website": "个人主页",
}


def _unwrap_resume(value: dict | None) -> dict:
    source = value if isinstance(value, dict) else {}
    for key in ("resume", "content", "optimized_resume"):
        if isinstance(source.get(key), dict):
            return source[key]
    return source


def _first_value(source: dict, aliases: tuple[str, ...]):
    for key in aliases:
        value = source.get(key)
        if value not in (None, "", []):
            return value
    return None


def _flatten_lines(value) -> list[str]:
    if value in (None, "", []):
        return []
    if isinstance(value, list):
        result: list[str] = []
        for item in value:
            result.extend(_flatten_lines(item))
        return result
    if isinstance(value, dict):
        result = []
        for nested in value.values():
            result.extend(_flatten_lines(nested))
        return result
    text = str(value).replace("\r", "\n")
    parts = re.split(r"\n+|\s*[●•▪◦]\s*", text)
    return [part.strip().lstrip("-–— ") for part in parts if part.strip().lstrip("-–— ")]


def _unique_lines(values: list[str]) -> list[str]:
    return list(dict.fromkeys(value.strip() for value in values if value and value.strip()))


def _plain_text_lines(value) -> list[str]:
    return _unique_lines(_flatten_lines(value))


def _resume_search_text(resume: dict | None) -> str:
    content = _unwrap_resume(resume)
    return " ".join(_plain_text_lines(content)).lower()


def _jd_terms(jd: dict | None) -> list[str]:
    source = jd if isinstance(jd, dict) else {}
    values: list[str] = []
    for key in ("keywords", "requirements", "responsibilities", "preferred", "title"):
        values.extend(_plain_text_lines(source.get(key)))
    terms: list[str] = []
    for value in values:
        terms.extend(re.split(r"[、,，;；|/\s]+", value))
    return _unique_lines([term for term in terms if 2 <= len(term) <= 18])[:24]


def _score_from_ratio(ratio: float, floor: int = 58, ceiling: int = 96) -> int:
    ratio = max(0.0, min(1.0, ratio))
    return round(floor + (ceiling - floor) * ratio)


def _star_rating(score: int) -> str:
    filled = max(1, min(5, round(score / 20)))
    return "★" * filled + "☆" * (5 - filled)


def _infer_role_family(text: str) -> str:
    lowered = text.lower()
    if any(word in lowered for word in ("销售", "客户", "商务", "渠道")):
        return "销售管理方向"
    if any(word in lowered for word in ("产品", "用户", "需求", "原型")):
        return "产品运营方向"
    if any(word in lowered for word in ("运营", "增长", "内容", "社群")):
        return "运营增长方向"
    if any(word in lowered for word in ("财务", "会计", "审计", "税务")):
        return "财务管理方向"
    if any(word in lowered for word in ("工程", "生产", "质量", "设备", "制造")):
        return "工程运营方向"
    if any(word in lowered for word in ("设计", "品牌", "视觉", "创意")):
        return "品牌创意方向"
    if any(word in lowered for word in ("开发", "技术", "系统", "数据", "算法")):
        return "技术专业方向"
    return "通用职业发展方向"


def _recommended_roles(role_family: str, title: str = "") -> list[str]:
    if "销售" in role_family:
        return ["销售主管", "商务经理", "渠道经理"]
    if "产品" in role_family:
        return ["产品经理", "项目经理", "产品运营"]
    if "运营" in role_family:
        return ["运营经理", "增长运营", "用户运营"]
    if "财务" in role_family:
        return ["财务主管", "会计主管", "审计专员"]
    if "工程" in role_family:
        return ["生产主管", "质量工程师", "项目工程师"]
    if "品牌" in role_family:
        return ["品牌策划", "视觉设计师", "新媒体运营"]
    if "技术" in role_family:
        return ["软件工程师", "数据分析师", "技术支持工程师"]
    return _unique_lines([title, "岗位专员", "项目执行", "业务主管"])[:3]


def build_career_profile(resumes: list[dict] | None, generations: list[dict] | None = None) -> dict:
    """Build a lightweight career asset profile from the user's existing resume records."""
    items = resumes or []
    completed_generations = [item for item in generations or [] if item.get("status") == "completed"]
    resume = next((item for item in items if item.get("is_default")), items[0] if items else {})
    content = _unwrap_resume(resume.get("content") if isinstance(resume, dict) else {})
    text = _resume_search_text(content)
    title = str(content.get("title") or "").strip()
    role_family = _infer_role_family(" ".join([title, text]))
    skills = _plain_text_lines(content.get("skills"))[:6]
    experience = content.get("experience") if isinstance(content.get("experience"), list) else []
    projects = content.get("projects") if isinstance(content.get("projects"), list) else []
    strength_pool = []
    for line in skills:
        strength_pool.append(f"具备{line.rstrip('。')}相关经验。")
    if experience:
        strength_pool.append("已有可沉淀的工作经历，可继续提炼职责、成果和岗位关键词。")
    if projects:
        strength_pool.append("包含项目经历，适合进一步包装成解决问题和推动结果的案例。")
    if not strength_pool and text:
        strength_pool.append("职业经历已建立，下一步可补充更清晰的职责和成果描述。")
    score = min(
        95,
        45
        + min(len(experience), 4) * 8
        + min(len(projects), 3) * 5
        + min(len(skills), 5) * 3
        + min(len(completed_generations), 4) * 2,
    )
    if not text:
        score = 0
    capability_names = skills[:4] or [role_family.replace("方向", ""), "岗位理解", "经历表达"]
    capabilities = [
        {
            "name": name[:18],
            "score": max(62, min(95, score - index * 4)),
            "stars": _star_rating(max(62, score - index * 4)),
        }
        for index, name in enumerate(capability_names[:4])
    ]
    actions = []
    if not items:
        actions.append("先上传或填写一份基础简历，建立职业档案。")
    if items and not completed_generations:
        actions.append("选择一个目标岗位，生成第一份岗位专属简历。")
    if score and score < 82:
        actions.append("补充更多可验证的成果、规模和项目细节，提升竞争力评分。")
    if not actions:
        actions.append("继续用不同岗位测试匹配度，沉淀更完整的求职策略。")
    return {
        "status": "ready" if text else "empty",
        "direction": role_family,
        "title": title or role_family,
        "competitiveness": score,
        "resume_count": len(items),
        "generation_count": len(completed_generations),
        "capabilities": capabilities,
        "strengths": strength_pool[:4],
        "recommended_roles": _recommended_roles(role_family, title),
        "next_actions": actions[:4],
    }


def build_jd_insight(jd: dict | None, resume: dict | None = None) -> dict:
    source = jd if isinstance(jd, dict) else {}
    resume_text = _resume_search_text(resume)
    terms = _jd_terms(source)
    matched_terms = [term for term in terms if term.lower() in resume_text]
    match_ratio = len(matched_terms) / max(len(terms), 1)
    responsibilities = _plain_text_lines(source.get("responsibilities"))
    requirements = _plain_text_lines(source.get("requirements"))
    preferred = _plain_text_lines(source.get("preferred"))
    core = _unique_lines(requirements + responsibilities + preferred + terms)[:5]
    if not core and source.get("title"):
        core = [f"围绕{source.get('title')}岗位要求进行能力匹配。"]
    gaps = []
    for item in core:
        item_terms = [term for term in re.split(r"[、,，;；|/\s]+", item) if len(term) >= 2]
        if resume_text and any(term.lower() in resume_text for term in item_terms):
            continue
        gaps.append(item)
    score = _score_from_ratio(match_ratio, floor=55, ceiling=92) if resume_text else 0
    suggestions = []
    if gaps:
        suggestions.append(f"重点补强：{gaps[0][:48]}。")
    if matched_terms:
        suggestions.append(f"生成简历时优先突出：{'、'.join(matched_terms[:5])}。")
    suggestions.append("如果有真实业绩、团队规模、客户数量或增长结果，建议补充到基础简历里。")
    return {
        "title": source.get("title") or "目标岗位",
        "company": source.get("company") or "",
        "core_requirements": core,
        "match_score": score,
        "match_level": _star_rating(score) if score else "待匹配",
        "advantages": matched_terms[:6],
        "gaps": gaps[:5],
        "suggestions": suggestions[:4],
    }


def score_generated_resume(resume: dict | None, jd: dict | None, design: dict | None = None) -> dict:
    content = _unwrap_resume(resume)
    terms = _jd_terms(jd)
    resume_text = _resume_search_text(content)
    matched_terms = [term for term in terms if term.lower() in resume_text]
    keyword_coverage = _score_from_ratio(len(matched_terms) / max(len(terms), 1), floor=54, ceiling=96)
    experience = content.get("experience") if isinstance(content.get("experience"), list) else []
    projects = content.get("projects") if isinstance(content.get("projects"), list) else []
    detail_count = sum(len(item.get("details") or []) for item in experience + projects if isinstance(item, dict))
    content_quality = min(95, 62 + min(detail_count, 12) * 2 + (8 if content.get("summary") else 0))
    job_match = round((keyword_coverage * 0.7) + (content_quality * 0.3))
    theme_id = (design or {}).get("theme_id", "auto")
    visual = 88 if theme_id == "ats_mono" else 93
    overall = round(content_quality * 0.35 + job_match * 0.35 + keyword_coverage * 0.2 + visual * 0.1)
    highlights = []
    if matched_terms:
        highlights.append(f"覆盖岗位关键词：{'、'.join(matched_terms[:6])}。")
    if content.get("summary"):
        highlights.append("已生成“经验与能力概述”，方便招聘方快速判断匹配度。")
    if detail_count:
        highlights.append("已将经历拆成更适合阅读的职责与成果描述。")
    if (design or {}).get("label"):
        highlights.append(f"视觉风格已匹配为：{(design or {}).get('label')}。")
    checks = [
        {"name": "AI痕迹", "status": "通过", "detail": "已限制空话套话，优先保留真实经历。"},
        {"name": "真实性", "status": "通过", "detail": "系统要求不新增公司、岗位、学历、时间和虚构数字。"},
        {
            "name": "岗位匹配",
            "status": "通过" if keyword_coverage >= 70 else "需补充",
            "detail": f"关键词覆盖评分 {keyword_coverage} 分。",
        },
        {"name": "排版", "status": "通过", "detail": "已输出 Word 与 PDF 两种文件，支持真实预览。"},
    ]
    return {
        "overall": overall,
        "dimensions": {
            "content_quality": content_quality,
            "job_match": job_match,
            "keyword_coverage": keyword_coverage,
            "visual_professionalism": visual,
        },
        "matched_keywords": matched_terms[:10],
        "highlights": highlights[:6],
        "checks": checks,
    }


def _normalize_entry(value, kind: str) -> dict:
    if not isinstance(value, dict):
        return {"details": _unique_lines(_flatten_lines(value))}
    date = _first_value(value, ("period", "date", "date_range", "duration", "time", "时间", "日期"))
    if kind == "work":
        primary_key = "company"
        primary = _first_value(value, ("company", "organization", "employer", "公司", "单位"))
        role = _first_value(value, ("role", "position", "job_title", "title", "岗位", "职位"))
    elif kind == "project":
        primary_key = "project"
        primary = _first_value(value, ("project", "project_name", "name", "项目", "项目名称"))
        role = _first_value(value, ("role", "position", "job_title", "title", "角色", "职责"))
    else:
        primary_key = "school"
        primary = _first_value(value, ("school", "institution", "university", "学校", "院校"))
        degree = _first_value(value, ("degree", "qualification", "学历"))
        major = _first_value(value, ("major", "专业"))
        role = " · ".join(str(item) for item in (degree, major) if item not in (None, ""))
    details: list[str] = []
    for key in ("details", "description", "responsibilities", "bullets", "content", "详情", "说明", "职责"):
        details.extend(_flatten_lines(value.get(key)))
    for key, prefix in (("achievements", "成果"), ("highlights", "亮点"), ("业绩", "成果"), ("成果", "成果")):
        details.extend(f"{prefix}：{line}" for line in _flatten_lines(value.get(key)))
    normalized = {"period": str(date or ""), primary_key: str(primary or ""), "role": str(role or "")}
    normalized["details"] = _unique_lines(details)
    if not any(normalized.values()):
        normalized["details"] = _unique_lines(_flatten_lines(value))
    return normalized


def _as_list(value) -> list:
    if value in (None, "", []):
        return []
    return value if isinstance(value, list) else [value]


def _match_text(value) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _entry_match_score(candidate: dict, fallback: dict, primary_key: str) -> int:
    score = 0
    for key, exact_score, partial_score in ((primary_key, 8, 4), ("role", 5, 2)):
        left = _match_text(candidate.get(key))
        right = _match_text(fallback.get(key))
        if not left or not right:
            continue
        if left == right:
            score += exact_score
        elif left in right or right in left:
            score += partial_score
    return score


def _preserve_missing_entry_periods(entries: list[dict], fallback_entries: list[dict], kind: str) -> list[dict]:
    if not fallback_entries:
        return entries
    primary_key = {"work": "company", "project": "project", "education": "school"}[kind]
    used_fallback_indexes: set[int] = set()
    merged: list[dict] = []
    for index, entry in enumerate(entries):
        if str(entry.get("period") or "").strip():
            merged.append(entry)
            continue
        best_index: int | None = None
        best_score = 0
        for fallback_index, fallback_entry in enumerate(fallback_entries):
            if fallback_index in used_fallback_indexes or not str(fallback_entry.get("period") or "").strip():
                continue
            score = _entry_match_score(entry, fallback_entry, primary_key)
            if score > best_score:
                best_index = fallback_index
                best_score = score
        if best_index is None and index < len(fallback_entries):
            fallback_entry = fallback_entries[index]
            if str(fallback_entry.get("period") or "").strip():
                best_index = index
        if best_index is not None:
            entry = {**entry, "period": fallback_entries[best_index]["period"]}
            used_fallback_indexes.add(best_index)
        merged.append(entry)
    return merged


def _period_sort_key(entry: dict) -> tuple[int, int, int]:
    period = str(entry.get("period") or "").strip().lower()
    if not period:
        return (0, 0, 0)
    dates = [
        (int(year), int(month or 1))
        for year, month in re.findall(r"(?<!\d)((?:19|20)\d{2})(?:[./年-]\s*(\d{1,2}))?", period)
    ]
    if not dates:
        return (0, 0, 0)
    start_year, start_month = dates[0]
    if re.search(r"至今|现在|目前|present|current|now", period):
        end_year, end_month = 9999, 12
    else:
        end_year, end_month = dates[-1]
    return (1, end_year * 100 + end_month, start_year * 100 + start_month)


def normalize_resume_content(candidate: dict | None, original: dict | None = None) -> dict:
    source = _unwrap_resume(candidate)
    fallback = _unwrap_resume(original)

    def pick(aliases: tuple[str, ...], default):
        return _first_value(source, aliases) or _first_value(fallback, aliases) or default

    name = _first_value(fallback, ("name", "full_name", "姓名")) or pick(("name", "full_name", "姓名"), "")
    title = pick(("title", "target", "objective", "job_target", "position", "求职方向"), "")
    contact = pick(("contact", "contacts", "contact_info", "联系方式"), {})
    if not isinstance(contact, dict):
        contact = {"contact": str(contact)} if contact else {}
    for target, aliases in (
        ("phone", ("phone", "mobile", "telephone", "tel", "电话", "手机", "联系电话")),
        ("email", ("email", "mail", "邮箱", "电子邮箱")),
        ("age", ("age", "birth_age", "年龄", "年纪")),
        ("birth_date", ("birth_date", "birthday", "出生年月", "出生日期", "生日")),
        ("birth_year", ("birth_year", "出生年份")),
        ("gender", ("gender", "sex", "性别")),
        ("location", ("location", "city", "address", "所在地", "现居地", "地址", "城市")),
    ):
        if not contact.get(target):
            for alias in aliases:
                if contact.get(alias):
                    contact[target] = contact[alias]
                    break
    for target, aliases in (
        ("phone", ("phone", "mobile", "telephone", "tel", "电话", "手机", "联系电话")),
        ("email", ("email", "mail", "邮箱", "电子邮箱")),
        ("age", ("age", "birth_age", "年龄", "年纪")),
        ("birth_date", ("birth_date", "birthday", "出生年月", "出生日期", "生日")),
        ("birth_year", ("birth_year", "出生年份")),
        ("gender", ("gender", "sex", "性别")),
        ("location", ("location", "city", "address", "所在地", "现居地", "地址", "城市")),
    ):
        value = pick(aliases, "")
        if value and not contact.get(target):
            contact[target] = value
    fallback_contact = fallback.get("contact") if isinstance(fallback.get("contact"), dict) else {}
    contact = {**contact, **{key: value for key, value in fallback_contact.items() if value}}
    for target, aliases in (
        ("phone", ("phone", "mobile", "telephone", "tel", "电话", "手机", "联系电话")),
        ("email", ("email", "mail", "邮箱", "电子邮箱")),
        ("age", ("age", "birth_age", "年龄", "年纪")),
        ("birth_date", ("birth_date", "birthday", "出生年月", "出生日期", "生日")),
        ("birth_year", ("birth_year", "出生年份")),
        ("gender", ("gender", "sex", "性别")),
        ("location", ("location", "city", "address", "所在地", "现居地", "地址", "城市")),
    ):
        if not contact.get(target):
            for alias in aliases:
                if contact.get(alias):
                    contact[target] = contact[alias]
                    break
    for alias in (
        "mobile",
        "telephone",
        "tel",
        "电话",
        "手机",
        "联系电话",
        "mail",
        "邮箱",
        "电子邮箱",
        "birth_age",
        "年龄",
        "年纪",
        "birthday",
        "出生年月",
        "出生日期",
        "生日",
        "sex",
        "性别",
        "city",
        "address",
        "所在地",
        "现居地",
        "地址",
        "城市",
    ):
        contact.pop(alias, None)
    summary = pick(("summary", "profile", "about", "introduction", "self_evaluation", "个人概述"), "")
    skills = _unique_lines(_flatten_lines(pick(("skills", "professional_skills", "core_skills", "专业技能"), [])))
    certificates = _unique_lines(_flatten_lines(pick(("certificates", "certifications", "awards", "证书"), [])))
    experience_aliases = ("experience", "work_experience", "employment", "work_history", "工作经历")
    project_aliases = ("projects", "project_experience", "项目经历")
    education_aliases = ("education", "education_experience", "academic_background", "教育经历")
    raw_experience = pick(experience_aliases, [])
    raw_projects = pick(project_aliases, [])
    raw_education = pick(education_aliases, [])
    fallback_experience = _first_value(fallback, experience_aliases) or []
    fallback_projects = _first_value(fallback, project_aliases) or []
    fallback_education = _first_value(fallback, education_aliases) or []
    experience_items = _as_list(raw_experience)
    project_items = _as_list(raw_projects)
    education_items = _as_list(raw_education)
    fallback_experience_items = [_normalize_entry(item, "work") for item in _as_list(fallback_experience)]
    fallback_project_items = [_normalize_entry(item, "project") for item in _as_list(fallback_projects)]
    fallback_education_items = [_normalize_entry(item, "education") for item in _as_list(fallback_education)]
    experience = sorted(
        _preserve_missing_entry_periods(
            [_normalize_entry(item, "work") for item in experience_items],
            fallback_experience_items,
            "work",
        ),
        key=_period_sort_key,
        reverse=True,
    )
    projects = _preserve_missing_entry_periods(
        [_normalize_entry(item, "project") for item in project_items],
        fallback_project_items,
        "project",
    )
    education = _preserve_missing_entry_periods(
        [_normalize_entry(item, "education") for item in education_items],
        fallback_education_items,
        "education",
    )
    return {
        "name": str(name or ""),
        "title": str(title or ""),
        "contact": {str(key): str(value) for key, value in contact.items() if value not in (None, "")},
        "summary": str(summary or ""),
        "skills": skills,
        "experience": [item for item in experience if any(item.values())],
        "projects": [item for item in projects if any(item.values())],
        "education": [item for item in education if any(item.values())],
        "certificates": certificates,
    }


def _avatar_candidate(data: bytes, max_dimension: int | None = None) -> tuple[bytes, float] | None:
    try:
        with Image.open(BytesIO(data)) as raw_image:
            image = ImageOps.exif_transpose(raw_image)
            width, height = image.size
            if width < 80 or height < 80 or (max_dimension and max(width, height) > max_dimension):
                return None
            ratio = width / height
            if not 0.5 <= ratio <= 1.2:
                return None
            if image.mode not in {"RGB", "L"}:
                background = Image.new("RGB", image.size, "white")
                alpha = image.getchannel("A") if "A" in image.getbands() else None
                background.paste(image.convert("RGB"), mask=alpha)
                image = background
            else:
                image = image.convert("RGB")
            image.thumbnail((600, 800), Image.Resampling.LANCZOS)
            output = BytesIO()
            image.save(output, format="JPEG", quality=90, optimize=True)
            score = (width * height) / (1 + abs(ratio - 0.78) * 3)
            return output.getvalue(), score
    except (OSError, ValueError):
        return None


def prepare_avatar_image(data: bytes) -> bytes:
    candidate = _avatar_candidate(data)
    if candidate is None:
        raise ValueError("头像图片需清晰完整，建议使用正方形或竖版 JPG/PNG/WebP")
    return candidate[0]


def _avatar_layout_size(data: bytes, max_width_mm: float = 32, max_height_mm: float = 30) -> tuple[float, float]:
    with Image.open(BytesIO(data)) as image:
        width, height = image.size
    scale = min(max_width_mm / width, max_height_mm / height)
    return width * scale, height * scale


def extract_resume_avatar(filename: str, data: bytes) -> bytes | None:
    suffix = filename.lower().rsplit(".", 1)[-1]
    candidates: list[tuple[bytes, float]] = []
    if suffix == "docx":
        try:
            with zipfile.ZipFile(BytesIO(data)) as archive:
                for name in archive.namelist():
                    if name.startswith("word/media/"):
                        candidate = _avatar_candidate(archive.read(name))
                        if candidate:
                            candidates.append(candidate)
        except (OSError, zipfile.BadZipFile):
            return None
    elif suffix == "pdf":
        try:
            reader = PdfReader(BytesIO(data))
            for page in reader.pages[:3]:
                for image_file in page.images:
                    candidate = _avatar_candidate(image_file.data, max_dimension=1200)
                    if candidate:
                        candidates.append(candidate)
        except Exception:
            return None
    return max(candidates, key=lambda item: item[1])[0] if candidates else None


def _set_docx_font(run, size: float, bold: bool = False, color: str = "27324A") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _remove_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def _shade_docx_cell(cell, color: str) -> None:
    """Apply a solid Word cell background without relying on a document theme."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), str(color or "FFFFFF").lstrip("#"))


def _add_paragraph_bottom_border(paragraph, color: str = "8AA4D6") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _contact_text(contact: dict) -> str:
    parts = []
    for key, value in contact.items():
        if value:
            parts.append(f"{CONTACT_LABELS.get(str(key).lower(), str(key))}：{value}")
    return "  ·  ".join(parts)


def _skills_summary(skills: list[str]) -> str:
    items = [str(item).strip() for item in skills if str(item).strip()]
    if not items:
        return ""
    if len(items) <= 2 and any(len(item) >= 16 or item.endswith(("。", "！", "；")) for item in items):
        return " ".join(item if item.endswith(("。", "！", "；")) else f"{item}。" for item in items)
    return f"具备{'、'.join(items[:5])}等与目标岗位相关的经验。"


def _experience_capability_overview(resume: dict) -> str:
    summary = str(resume.get("summary") or "").strip()
    return summary or _skills_summary(resume.get("skills") or [])


def _docx_section_title(document: Document, title: str, color: str = "244FB8") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    _set_docx_font(paragraph.add_run(title), 11, bold=True, color=color)
    _add_paragraph_bottom_border(paragraph, color=color)


def _entry_primary(entry: dict, kind: str) -> str:
    key = {"work": "company", "project": "project", "education": "school"}[kind]
    return str(entry.get(key) or entry.get("role") or "")


def _entry_heading(entry: dict, kind: str) -> str:
    primary = _entry_primary(entry, kind)
    role = str(entry.get("role") or "")
    if kind == "work" and role and role != primary:
        return f"{primary}  {role}".strip()
    return primary


def _docx_entry(document: Document, entry: dict, kind: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Mm(135)
    table.columns[1].width = Mm(35)
    _remove_table_borders(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    left_p = left.paragraphs[0]
    left_p.paragraph_format.keep_with_next = True
    _set_docx_font(left_p.add_run(_entry_heading(entry, kind)), 10.4, bold=True)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.keep_with_next = True
    _set_docx_font(right_p.add_run(str(entry.get("period") or "")), 8.5, color="69738A")
    if kind != "work" and entry.get("role") and str(entry.get("role")) != _entry_primary(entry, kind):
        role_paragraph = document.add_paragraph()
        role_paragraph.paragraph_format.space_after = Pt(2)
        role_paragraph.paragraph_format.keep_with_next = True
        _set_docx_font(role_paragraph.add_run(str(entry["role"])), 8.7, bold=True, color="43516D")
    for detail in entry.get("details") or []:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Mm(6)
        paragraph.paragraph_format.first_line_indent = Mm(-2.5)
        paragraph.paragraph_format.space_after = Pt(1)
        _set_docx_font(paragraph.add_run(str(detail)), 8.6)


def _docx_write_body_sections(document: Document, resume: dict, primary: str, *, include_overview: bool = True, include_certificates: bool = True) -> None:
    overview = _experience_capability_overview(resume)
    if include_overview and overview:
        _docx_section_title(document, "经验与能力概述", primary)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.25
        _set_docx_font(paragraph.add_run(overview), 8.8, color="33415C")
    for title, key, kind in (
        ("工作经历", "experience", "work"),
        ("项目经历", "projects", "project"),
        ("教育经历", "education", "education"),
    ):
        if resume.get(key):
            _docx_section_title(document, title, primary)
            for entry in resume[key]:
                _docx_entry(document, entry, kind)
    if include_certificates and resume.get("certificates"):
        _docx_section_title(document, "证书与荣誉", primary)
        paragraph = document.add_paragraph()
        _set_docx_font(paragraph.add_run("  ·  ".join(resume["certificates"])), 8.8)


def resume_to_docx(resume: dict, avatar_data: bytes | None = None, design: dict | None = None) -> bytes:
    resume = normalize_resume_content(resume)
    design = design or {}
    style = _lapiscv_theme(design)
    primary = str(style.get("accent") or "#173A70").lstrip("#")
    secondary = str(style.get("text") or "#43516D").lstrip("#")
    layout_variant = str(design.get("word_layout") or design.get("layout_variant") or "top_profile")
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(12 if layout_variant == "left_sidebar" else 14)
    section.right_margin = Mm(12 if layout_variant == "right_sidebar" else 14)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(8.4 if layout_variant in {"left_sidebar", "campus_compact"} else 8.6)
    normal.paragraph_format.space_after = Pt(2)

    # True structural differences (not only a colour strip).
    if layout_variant == "left_sidebar":
        shell = document.add_table(rows=1, cols=2)
        _remove_table_borders(shell)
        shell.columns[0].width = Mm(52)
        shell.columns[1].width = Mm(130)
        side, main = shell.cell(0, 0), shell.cell(0, 1)
        _shade_docx_cell(side, primary)
        side_p = side.paragraphs[0]
        _set_docx_font(side_p.add_run(resume.get("name") or "个人简历"), 14, bold=True, color="FFFFFF")
        if resume.get("title"):
            tp = side.add_paragraph()
            _set_docx_font(tp.add_run(str(resume["title"])), 8.5, bold=True, color="DCE8FF")
        contact = _contact_text(resume.get("contact") or {})
        if contact:
            cp = side.add_paragraph()
            _set_docx_font(cp.add_run(contact), 7.5, color="F2F6FF")
        overview = _experience_capability_overview(resume)
        if overview:
            op = side.add_paragraph()
            _set_docx_font(op.add_run("能力概述"), 8.5, bold=True, color="FFFFFF")
            op2 = side.add_paragraph()
            _set_docx_font(op2.add_run(overview), 7.5, color="F2F6FF")
        if resume.get("certificates"):
            cp0 = side.add_paragraph()
            _set_docx_font(cp0.add_run("证书"), 8.5, bold=True, color="FFFFFF")
            cp1 = side.add_paragraph()
            _set_docx_font(cp1.add_run(" · ".join(resume["certificates"])), 7.5, color="F2F6FF")
        # Write main column body into the right cell.
        first = True
        for title, key, kind in (
            ("工作经历", "experience", "work"),
            ("项目经历", "projects", "project"),
            ("教育经历", "education", "education"),
        ):
            if not resume.get(key):
                continue
            hp = main.paragraphs[0] if first and not main.paragraphs[0].text else main.add_paragraph()
            first = False
            _set_docx_font(hp.add_run(title), 11, bold=True, color=primary)
            _add_paragraph_bottom_border(hp, color=primary)
            for entry in resume[key]:
                _docx_entry_into_cell(main, entry, kind, primary)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    if layout_variant in {"banner_timeline", "creative_asymmetry"}:
        banner = document.add_table(rows=1, cols=1)
        _remove_table_borders(banner)
        banner.cell(0, 0).width = Mm(182)
        _shade_docx_cell(banner.cell(0, 0), primary)
        cell = banner.cell(0, 0)
        name_p = cell.paragraphs[0]
        _set_docx_font(name_p.add_run(resume.get("name") or "个人简历"), 18, bold=True, color="FFFFFF")
        if resume.get("title"):
            tp = cell.add_paragraph()
            _set_docx_font(tp.add_run(str(resume["title"])), 9.5, bold=True, color="EEF4FF")
        contact = _contact_text(resume.get("contact") or {})
        if contact:
            cp = cell.add_paragraph()
            _set_docx_font(cp.add_run(contact), 8, color="E8EEF8")
    else:
        header_cols = 2 if avatar_data else 1
        header = document.add_table(rows=1, cols=header_cols)
        header.alignment = WD_TABLE_ALIGNMENT.LEFT
        header.autofit = False
        _remove_table_borders(header)
        left_cell = header.cell(0, 0)
        name_paragraph = left_cell.paragraphs[0]
        name_paragraph.paragraph_format.space_after = Pt(2)
        name_color = primary
        if layout_variant == "campus_compact":
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_docx_font(name_paragraph.add_run(resume.get("name") or "个人简历"), 18 if layout_variant == "executive_minimal" else 20, bold=True, color=name_color)
        if resume.get("title"):
            title_paragraph = left_cell.add_paragraph()
            title_paragraph.paragraph_format.space_after = Pt(3)
            if layout_variant == "campus_compact":
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_docx_font(title_paragraph.add_run(str(resume["title"])), 9.8, bold=True, color=secondary)
        contact = _contact_text(resume.get("contact") or {})
        if contact:
            contact_paragraph = left_cell.add_paragraph()
            if layout_variant == "campus_compact":
                contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_docx_font(contact_paragraph.add_run(contact), 8.2, color="69738A")
        if avatar_data:
            avatar_cell = header.cell(0, 1)
            avatar_cell.width = Mm(36)
            avatar_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            avatar_paragraph = avatar_cell.paragraphs[0]
            avatar_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            avatar_width, avatar_height = _avatar_layout_size(avatar_data)
            avatar_paragraph.add_run().add_picture(
                BytesIO(avatar_data), width=Mm(avatar_width), height=Mm(avatar_height)
            )
        separator = document.add_paragraph()
        separator.paragraph_format.space_after = Pt(1)
        _add_paragraph_bottom_border(separator, color=primary)

    if layout_variant == "split_columns":
        split = document.add_table(rows=1, cols=2)
        _remove_table_borders(split)
        left, right = split.cell(0, 0), split.cell(0, 1)
        overview = _experience_capability_overview(resume)
        if overview:
            p = left.paragraphs[0]
            _set_docx_font(p.add_run("经验与能力概述"), 11, bold=True, color=primary)
            p2 = left.add_paragraph()
            _set_docx_font(p2.add_run(overview), 8.6, color="33415C")
        if resume.get("education"):
            ep = left.add_paragraph()
            _set_docx_font(ep.add_run("教育经历"), 11, bold=True, color=primary)
            for entry in resume["education"]:
                _docx_entry_into_cell(left, entry, "education", primary)
        if resume.get("certificates"):
            cp = left.add_paragraph()
            _set_docx_font(cp.add_run("证书与荣誉"), 11, bold=True, color=primary)
            cp2 = left.add_paragraph()
            _set_docx_font(cp2.add_run(" · ".join(resume["certificates"])), 8.6)
        first_right = True
        for title, key, kind in (("工作经历", "experience", "work"), ("项目经历", "projects", "project")):
            if not resume.get(key):
                continue
            hp = right.paragraphs[0] if first_right and not right.paragraphs[0].text else right.add_paragraph()
            first_right = False
            _set_docx_font(hp.add_run(title), 11, bold=True, color=primary)
            for entry in resume[key]:
                _docx_entry_into_cell(right, entry, kind, primary)
    else:
        _docx_write_body_sections(
            document,
            resume,
            primary,
            include_overview=True,
            include_certificates=layout_variant != "left_sidebar",
        )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_entry_into_cell(cell, entry: dict, kind: str, primary: str) -> None:
    heading = _entry_heading(entry, kind)
    period = str(entry.get("period") or "")
    hp = cell.add_paragraph()
    _set_docx_font(hp.add_run(heading), 9.2, bold=True, color="27324A")
    if period:
        _set_docx_font(hp.add_run(f"  {period}"), 8.2, color=primary)
    for detail in entry.get("details") or entry.get("description") or []:
        line = str(detail).strip()
        if not line:
            continue
        dp = cell.add_paragraph()
        _set_docx_font(dp.add_run(f"• {line}"), 8.2, color="33415C")


@lru_cache(maxsize=1)
def _pdf_font_name() -> str:
    candidates = (
        Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
        Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    )
    for path in candidates:
        if path.is_file():
            try:
                pdfmetrics.registerFont(TTFont("ResumeCJK", str(path), subfontIndex=0))
                return "ResumeCJK"
            except TTFError:
                continue
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


RESUME_DESIGN_THEMES = {
    "tech_indigo": {
        "label": "技术专业", "primary": "#284C9B", "ribbon": "#284C9B", "soft": "#EAF0FF",
        "ink": "#202A42", "muted": "#5F6B84", "layout": "timeline_focus", "colored": True,
        "layout_variant": "left_sidebar", "template_id": "classic", "density": "compact",
    },
    "operations_terra": {
        "label": "工程运营", "primary": "#1F6268", "ribbon": "#1F6268", "soft": "#E7F2EF",
        "ink": "#24383B", "muted": "#5D6E70", "layout": "timeline_focus", "colored": True,
        "layout_variant": "banner_timeline", "template_id": "classic", "density": "compact",
    },
    "executive_navy": {
        "label": "稳重商务", "primary": "#263B59", "ribbon": "#263B59", "soft": "#EDF1F5",
        "ink": "#202B3B", "muted": "#667386", "layout": "accent_header", "colored": True,
        "layout_variant": "executive_minimal", "template_id": "classic", "density": "airy",
    },
    "care_teal": {
        "label": "教育医疗", "primary": "#16756F", "ribbon": "#16756F", "soft": "#E7F5F1",
        "ink": "#1E3938", "muted": "#607A78", "layout": "timeline_focus", "colored": True,
        "layout_variant": "banner_timeline", "template_id": "serif", "density": "balanced",
    },
    "creative_plum": {
        "label": "品牌创意", "primary": "#75416F", "ribbon": "#75416F", "soft": "#F7EDF4",
        "ink": "#3B293C", "muted": "#786578", "layout": "timeline_focus", "colored": True,
        "layout_variant": "creative_asymmetry", "template_id": "serif", "density": "balanced",
    },
    "ats_mono": {
        "label": "打印友好", "primary": "#3D4654", "ribbon": "#697484", "soft": "#F4F5F7",
        "ink": "#28313F", "muted": "#687281", "layout": "accent_header", "colored": False,
        "layout_variant": "top_profile", "template_id": "classic", "density": "balanced",
    },
}


def _jd_design_text(jd: dict) -> str:
    values = [jd.get(key) for key in ("title", "company", "responsibilities", "requirements", "keywords", "raw_text")]
    return " ".join(str(value) for value in values if value).lower()


def fallback_resume_design(jd: dict, requested_theme: str = "auto") -> dict:
    if requested_theme in RESUME_DESIGN_THEMES:
        theme_id = requested_theme
        reason = "已按用户选择应用主题"
    else:
        text = _jd_design_text(jd)
        groups = (
            ("creative_plum", "设计|品牌|市场|营销|新媒体|广告|内容运营"),
            ("care_teal", "教育|教师|学校|医疗|医药|护士|医生|心理"),
            ("operations_terra", "制造|生产|运维|工程|施工|设备|物流|建筑|质量"),
            ("executive_navy", "金融|银行|财务|审计|咨询|法务|行政|人力|政务"),
            ("tech_indigo", "软件|开发|数据|算法|产品经理|互联网|人工智能|ai|it|测试"),
        )
        theme_id = next((theme for theme, pattern in groups if re.search(pattern, text)), "executive_navy")
        reason = "根据岗位信息自动匹配"
    theme = RESUME_DESIGN_THEMES[theme_id]
    return {
        "theme_id": theme_id,
        "label": theme["label"],
        "layout_id": theme["layout"],
        "template_id": theme.get("template_id") or ("classic" if theme_id not in {"care_teal", "creative_plum"} else "serif"),
        "density": theme.get("density") or "balanced",
        # Always pin a concrete layout family so “auto” is not a monochrome clone.
        "layout_variant": theme.get("layout_variant") or "top_profile",
        "word_layout": theme.get("layout_variant") or "top_profile",
        "header_mode": "side" if theme.get("layout_variant") == "left_sidebar" else "banner" if "banner" in str(theme.get("layout_variant")) else "compact",
        "section_style": "block",
        "avatar_mode": "circle" if theme_id in {"care_teal", "creative_plum"} else "square",
        "role_family": "",
        "company_type": "",
        "reason": reason,
        "source": "manual" if requested_theme in RESUME_DESIGN_THEMES else "fallback",
    }


def resolve_resume_design(jd: dict, requested_theme: str = "auto", planned: dict | None = None) -> dict:
    design = fallback_resume_design(jd, requested_theme)
    if requested_theme != "auto" or not isinstance(planned, dict):
        return design
    theme_id = str(planned.get("theme_id") or "")
    if theme_id not in RESUME_DESIGN_THEMES or theme_id == "ats_mono":
        return design
    theme = RESUME_DESIGN_THEMES[theme_id]
    base = fallback_resume_design(jd, theme_id)
    layout_id = str(planned.get("layout_id") or "")
    template_id = str(planned.get("template_id") or "")
    density = str(planned.get("density") or "")
    variant = str(planned.get("layout_variant") or base.get("layout_variant") or "top_profile")
    allowed_variants = {
        "top_profile",
        "left_sidebar",
        "right_sidebar",
        "banner_timeline",
        "split_columns",
        "executive_minimal",
        "creative_asymmetry",
        "campus_compact",
    }
    if variant not in allowed_variants:
        variant = base.get("layout_variant") or "top_profile"
    return {
        **base,
        "theme_id": theme_id,
        "label": theme["label"],
        "layout_id": layout_id if layout_id in {"accent_header", "timeline_focus"} else theme["layout"],
        "template_id": template_id if template_id in {"classic", "serif"} else base["template_id"],
        "density": density if density in {"compact", "balanced", "airy"} else base.get("density") or "balanced",
        "layout_variant": variant,
        "word_layout": variant,
        "role_family": str(planned.get("role_family") or "")[:80],
        "company_type": str(planned.get("company_type") or "")[:80],
        "reason": str(planned.get("reason") or design["reason"])[:80],
        "source": "ai",
    }


def apply_catalog_template(design: dict, template: dict | None) -> dict:
    """Apply a catalogued template's safe layout traits to the shared renderers."""
    if not template:
        return design
    result = dict(design)
    base_theme = str(template.get("base_theme") or "")
    if base_theme in RESUME_DESIGN_THEMES:
        result = fallback_resume_design({}, base_theme) | result
        result["theme_id"] = base_theme
    result.update({
        "label": str(template.get("name") or result.get("label") or "简历模板"),
        "layout_id": str(template.get("layout_id") or result.get("layout_id") or "accent_header"),
        "density": str(template.get("density") or result.get("density") or "balanced"),
        "catalog_template_id": str(template.get("id") or ""),
        "catalog_template_name": str(template.get("name") or ""),
        "catalog_template_version": int(template.get("version") or 1),
        "layout_variant": str(template.get("layout_variant") or "top_profile"),
        "header_mode": str(template.get("header_mode") or "compact"),
        "section_style": str(template.get("section_style") or "line"),
        "avatar_mode": str(template.get("avatar_mode") or "square"),
        "word_layout": str(template.get("word_layout") or template.get("layout_variant") or "top_profile"),
        "source": "catalog",
        "style_override": {
            "accent": str(template.get("accent") or ""),
            "soft": str(template.get("soft") or ""),
            "text": str(template.get("ink") or ""),
            "h1": str(template.get("ink") or ""),
            "template": "serif" if base_theme in {"care_teal", "creative_plum"} else "classic",
        },
    })
    return result


LAPISCV_ROOT = Path(__file__).resolve().parent / "templates" / "lapiscv"
LAPISCV_SECTION_ICONS = {
    "经验与能力概述": "&#xe782;",
    "工作经历": "&#xe618;",
    "项目经历": "&#xe635;",
    "教育经历": "&#xe80c;",
    "证书与荣誉": "&#xe673;",
}
LAPISCV_THEME_OVERRIDES = {
    "tech_indigo": {
        "template": "classic",
        "accent": "#284C9B",
        "text": "#202A42",
        "h1": "#202A42",
        "line_height": "1.74",
        "avatar_width": "27mm",
    },
    "operations_terra": {
        "template": "classic",
        "accent": "#1F6268",
        "text": "#24383B",
        "h1": "#24383B",
        "line_height": "1.72",
        "avatar_width": "27mm",
    },
    "executive_navy": {
        "template": "classic",
        "accent": "#263B59",
        "text": "#202B3B",
        "h1": "#202B3B",
        "line_height": "1.70",
        "avatar_width": "26mm",
    },
    "care_teal": {
        "template": "serif",
        "accent": "#16756F",
        "text": "#1E3938",
        "h1": "#1E3938",
        "line_height": "1.76",
        "avatar_width": "24mm",
    },
    "creative_plum": {
        "template": "serif",
        "accent": "#75416F",
        "text": "#3B293C",
        "h1": "#3B293C",
        "line_height": "1.76",
        "avatar_width": "24mm",
    },
    "ats_mono": {
        "template": "classic",
        "accent": "#3D4654",
        "text": "#28313F",
        "h1": "#28313F",
        "line_height": "1.66",
        "avatar_width": "25mm",
    },
}


def _html_escape(value: object) -> str:
    return (
        str(value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _html_paragraph(value: object) -> str:
    lines = [_html_escape(line.strip()) for line in str(value or "").splitlines() if line.strip()]
    return "<br>".join(lines)


def _lapiscv_theme(design: dict | None) -> dict:
    theme_id = str((design or {}).get("theme_id") or "executive_navy")
    theme = dict(LAPISCV_THEME_OVERRIDES.get(theme_id, LAPISCV_THEME_OVERRIDES["executive_navy"]))
    overrides = (design or {}).get("style_override") or {}
    if isinstance(overrides, dict):
        for key in ("accent", "text", "h1", "line_height", "avatar_width", "template"):
            value = overrides.get(key)
            if isinstance(value, str) and value.strip():
                theme[key] = value.strip()
    return theme


def _lapiscv_asset_uri(*parts: str) -> str:
    return (LAPISCV_ROOT.joinpath(*parts)).resolve().as_uri()


def _lapiscv_css(design: dict | None) -> str:
    theme = _lapiscv_theme(design)
    requested_template = str((design or {}).get("template_id") or "")
    template = requested_template if requested_template in {"classic", "serif"} else theme["template"]
    density = str((design or {}).get("density") or "balanced")
    density_vars = {
        "compact": {"text_size": "9.2pt", "line_height": "1.56", "h3_size": "10pt"},
        "balanced": {"text_size": "9.7pt", "line_height": theme["line_height"], "h3_size": "10.5pt"},
        "airy": {"text_size": "10pt", "line_height": "1.86", "h3_size": "10.8pt"},
    }.get(density, {"text_size": "9.7pt", "line_height": theme["line_height"], "h3_size": "10.5pt"})
    main_css = (LAPISCV_ROOT / "styles" / "main.css").read_text(encoding="utf-8")
    template_css = (LAPISCV_ROOT / "styles" / template / "vscode.css").read_text(encoding="utf-8")
    template_css = re.sub(r"@font-face\s*\{.*?\}", "", template_css, flags=re.DOTALL)
    return f"""
@page {{
  size: A4;
  margin: 13mm 15mm;
}}
@font-face {{
  font-family: "SourceHanSansCN";
  src: url("{_lapiscv_asset_uri("fonts", "SourceHanSansCN-Regular.ttf")}") format("truetype");
}}
@font-face {{
  font-family: "SourceHanSansCN";
  src: url("{_lapiscv_asset_uri("fonts", "SourceHanSansCN-Bold.ttf")}") format("truetype");
  font-weight: bold;
}}
@font-face {{
  font-family: "SourceHanSerifCN";
  src: url("{_lapiscv_asset_uri("fonts", "SourceHanSerifCN-Bold.ttf")}") format("truetype");
  font-weight: bold;
}}
@font-face {{
  font-family: "JetBrainsMono";
  src: url("{_lapiscv_asset_uri("fonts", "JetBrainsMono-Regular.ttf")}") format("truetype");
}}
@font-face {{
  font-family: "LapisCV Icon";
  src: url("{_lapiscv_asset_uri("fonts", "iconfont.ttf")}") format("truetype");
}}
{main_css}
{template_css}
html, body {{
  background: #fff !important;
}}
body {{
  width: auto;
  margin: 0;
  padding: 0;
  --color-accent: {theme["accent"]};
  --text-normal: {theme["text"]};
  --text-strong: {theme["text"]};
  --text-font: "WenQuanYi Zen Hei", "Microsoft YaHei", "SourceHanSansCN", sans-serif;
  --title-font: "WenQuanYi Zen Hei", "Microsoft YaHei", "SourceHanSansCN", sans-serif;
  --h1-color: {theme["h1"]};
  --h2-color: {theme["accent"]};
  --h3-color: {theme["text"]};
  --link-color: {theme["accent"]};
  --text-size: {density_vars["text_size"]};
  --line-height: {density_vars["line_height"]};
  --h3-size: {density_vars["h3_size"]};
  --avatar-width: {theme["avatar_width"]};
}}
h1 {{
  letter-spacing: .02em;
}}
h1,
h2,
h3,
p,
li,
blockquote,
code {{
  font-family: var(--text-font);
}}
h1,
h2,
h3 {{
  font-family: var(--title-font);
}}
h2 {{
  break-after: avoid;
}}
.lapiscv-header {{
  display: grid;
  grid-template-columns: minmax(0, 1fr) auto;
  gap: 8mm;
  align-items: start;
  margin-bottom: 5mm;
  break-inside: avoid;
}}
.lapiscv-header h1 {{
  margin: 0 0 1.2mm 0;
  padding: 0;
  text-align: left;
  line-height: 1.25;
}}
.lapiscv-title {{
  margin: 0 0 1.8mm 0;
  color: var(--color-accent);
  font-size: 10.5pt;
  line-height: 1.45;
}}
.lapiscv-contact {{
  margin: 0;
  color: var(--text-normal);
  font-size: 9pt;
  line-height: 1.55;
}}
.lapiscv-header .avatar {{
  display: block !important;
  float: none !important;
  position: static !important;
  top: auto !important;
  right: auto !important;
  width: var(--avatar-width) !important;
  height: var(--avatar-width) !important;
  margin: 0 !important;
  box-shadow: none !important;
}}
.lapiscv-header--no-avatar {{
  display: block;
  text-align: center;
}}
.lapiscv-header--no-avatar h1,
.lapiscv-header--no-avatar .lapiscv-title,
.lapiscv-header--no-avatar .lapiscv-contact {{
  text-align: center;
}}
.entry-title,
div[alt="entry-title"] {{
  gap: 8mm;
  break-inside: avoid;
  break-after: avoid;
}}
.entry-title h3,
div[alt="entry-title"] h3 {{
  flex: 1;
}}
.entry-title p,
div[alt="entry-title"] p {{
  white-space: nowrap;
  color: var(--color-accent);
}}
li {{
  break-inside: avoid;
}}
.lapiscv-entry {{
  break-inside: avoid-page;
  page-break-inside: avoid;
  margin-bottom: 1.4mm;
}}
.lapiscv-entry ul {{
  break-before: avoid;
}}
.lapiscv-section-summary {{
  margin-bottom: 1.5mm;
}}
.avatar,
img[alt="avatar"] {{
  object-fit: cover;
}}
/* Catalogue layout families: document structure changes, not just colours. */
.catalog-resume {{ width: 100%; --catalog-soft: {str((design or {}).get("style_override", {}).get("soft") or "#EEF3FF")}; }}
/* Strong structural fingerprints so layouts are obviously different in PDF. */
.catalog-banner {{ background: var(--color-accent); color: #fff; padding: 8mm 8mm 7mm; margin: 0 0 6mm; border-radius: 0 0 3mm 3mm; }}
.catalog-banner h1, .catalog-banner .lapiscv-contact, .catalog-banner .lapiscv-title {{ color: #fff !important; }}
.catalog-banner .lapiscv-title {{ opacity: .92; }}
.catalog-banner .avatar {{ width: 26mm !important; height: 26mm !important; border: 1.5mm solid #fff; border-radius: 50%; }}
.catalog-two-column {{ display: grid !important; grid-template-columns: 34% minmax(0, 1fr) !important; gap: 6mm; align-items: start; width: 100%; }}
.catalog-two-column--right {{ grid-template-columns: minmax(0, 1fr) 32% !important; }}
.catalog-sidebar {{ padding: 6mm 5mm; background: var(--catalog-soft); border-top: 4mm solid var(--color-accent); min-height: 240mm; border-radius: 2mm; }}
.catalog-sidebar--dark {{ background: var(--color-accent) !important; color: #fff; border-top-color: color-mix(in srgb, #000 18%, var(--color-accent)); }}
.catalog-sidebar--dark h1, .catalog-sidebar--dark h2, .catalog-sidebar--dark .lapiscv-contact, .catalog-sidebar--dark p, .catalog-sidebar--dark .lapiscv-title {{ color: #fff !important; }}
.catalog-sidebar .avatar {{ width: 28mm !important; height: 28mm !important; margin: 0 0 4mm !important; border: 1.2mm solid #fff; }}
.catalog-sidebar h1 {{ font-size: 16pt; margin: 0 0 2mm; line-height: 1.25; }}
.catalog-main {{ min-width: 0; }}
.catalog-main h2, .catalog-resume--split h2 {{ border-bottom: 1.2px solid color-mix(in srgb, var(--color-accent) 42%, #fff); padding-bottom: 1.5mm; margin-top: 3.5mm; }}
.catalog-resume--timeline {{ padding-left: 1mm; }}
.catalog-resume--timeline .lapiscv-entry {{ position: relative; padding-left: 7mm; margin-left: 2mm; border-left: 1.4px solid color-mix(in srgb, var(--color-accent) 55%, #fff); }}
.catalog-resume--timeline .lapiscv-entry::before {{ content: ""; position: absolute; left: -2.4mm; top: 2.2mm; width: 3.6mm; height: 3.6mm; border-radius: 50%; background: var(--color-accent); box-shadow: 0 0 0 1.2mm #fff; }}
.catalog-resume--executive {{ max-width: 168mm; margin: 0 auto; padding: 2mm 4mm; }}
.catalog-resume--executive .lapiscv-header {{ padding-bottom: 4.5mm; border-bottom: 2.2px solid var(--color-accent); margin-bottom: 4mm; }}
.catalog-resume--executive h1 {{ letter-spacing: .08em; font-size: 22pt; }}
.catalog-resume--executive h2 {{ letter-spacing: .18em; font-size: 10pt; text-transform: uppercase; border-bottom: none; color: var(--color-accent); }}
.catalog-resume--creative {{ border-top: 9mm solid var(--color-accent); padding-top: 5mm; }}
.catalog-resume--creative h2 {{ display: inline-block; padding: 1.6mm 5.5mm; background: var(--catalog-soft); border-left: 3.2mm solid var(--color-accent); border-radius: 0 2mm 2mm 0; }}
.catalog-resume--campus .lapiscv-header {{ text-align: center; display: block; border-bottom: 1.4px dashed var(--color-accent); padding-bottom: 4.5mm; margin-bottom: 4mm; }}
.catalog-resume--campus .lapiscv-header h1, .catalog-resume--campus .lapiscv-title, .catalog-resume--campus .lapiscv-contact {{ text-align: center; }}
.catalog-resume--campus .lapiscv-header .avatar {{ margin: 0 auto 3mm !important; border-radius: 3mm; display: block !important; }}
.catalog-resume--campus h2 {{ text-align: left; background: var(--catalog-soft); padding: 1.2mm 3mm; border-radius: 1.5mm; }}
.catalog-resume--split .catalog-summary-grid {{ display: grid !important; grid-template-columns: 1fr 1fr !important; gap: 7mm; margin-bottom: 4mm; }}
.catalog-resume--split .catalog-summary-grid > section {{ break-inside: avoid; padding: 3mm; background: var(--catalog-soft); border-radius: 2mm; }}
.catalog-resume--top .lapiscv-header {{ border-bottom: 1.5px solid color-mix(in srgb, var(--color-accent) 40%, #fff); padding-bottom: 3mm; margin-bottom: 4mm; }}
.catalog-resume--top h2 {{ border-left: 2.5mm solid var(--color-accent); padding-left: 3mm; }}
"""


def _lapiscv_avatar_src(avatar_data: bytes | None) -> str:
    if not avatar_data:
        return ""
    encoded = base64.b64encode(avatar_data).decode("ascii")
    return f"data:image/jpeg;base64,{encoded}"


def _lapiscv_contact(contact: dict) -> str:
    parts = []
    for key, value in (contact or {}).items():
        if not value:
            continue
        normalized = str(key).lower()
        label = CONTACT_LABELS.get(normalized, str(key))
        parts.append(f"{_html_escape(label)}：{_html_escape(value)}")
    return "　｜　".join(parts)


def _lapiscv_section(title: str) -> str:
    return f"<h2>{_html_escape(title)}</h2>"


def _lapiscv_details(entry: dict) -> list[str]:
    details = entry.get("details") or entry.get("description") or entry.get("bullets") or []
    if isinstance(details, str):
        details = [line.strip() for line in details.splitlines() if line.strip()]
    if not isinstance(details, list):
        return []
    return [str(item).strip() for item in details if str(item).strip()]


def _lapiscv_entry(entry: dict, kind: str) -> str:
    heading = _html_escape(_entry_heading(entry, kind))
    period = _html_escape(entry.get("period") or "")
    role = str(entry.get("role") or "")
    role_markup = ""
    if kind != "work" and role and role != _entry_primary(entry, kind):
        role_markup = f"<p><strong>{_html_escape(role)}</strong></p>"
    details = _lapiscv_details(entry)
    list_markup = ""
    if details:
        items = "".join(f"<li>{_html_paragraph(item)}</li>" for item in details)
        list_markup = f"<ul>{items}</ul>"
    return (
        '<section class="lapiscv-entry">'
        '<div class="entry-title">'
        f"<h3>{heading}</h3>"
        f"<p>{period}</p>"
        "</div>"
        f"{role_markup}{list_markup}"
        "</section>"
    )


def _catalog_header(resume: dict, avatar_src: str, *, banner: bool = False, compact: bool = False) -> str:
    contact_markup = _lapiscv_contact(resume.get("contact") or {})
    header_class = "lapiscv-header" if avatar_src else "lapiscv-header lapiscv-header--no-avatar"
    header_left = [f"<h1>{_html_escape(resume.get('name') or '个人简历')}</h1>"]
    if resume.get("title"):
        header_left.append(f'<p class="lapiscv-title"><strong>{_html_escape(resume["title"])}</strong></p>')
    if contact_markup:
        header_left.append(f'<p class="lapiscv-contact">{contact_markup}</p>')
    avatar_markup = f'<img class="avatar" alt="avatar" src="{avatar_src}">' if avatar_src else ""
    inner = f'<header class="{header_class}"><div>{"".join(header_left)}</div>{avatar_markup}</header>'
    return f'<div class="catalog-banner">{inner}</div>' if banner else inner


def _catalog_sections(resume: dict, *, include_overview: bool = True, include_certificates: bool = True) -> str:
    body: list[str] = []
    overview = _experience_capability_overview(resume)
    if include_overview and overview:
        body.append(_lapiscv_section("经验与能力概述"))
        body.append(f'<p class="lapiscv-section-summary">{_html_paragraph(overview)}</p>')
    for title, key, kind in (
        ("工作经历", "experience", "work"),
        ("项目经历", "projects", "project"),
        ("教育经历", "education", "education"),
    ):
        entries = resume.get(key) or []
        if entries:
            body.append(_lapiscv_section(title))
            body.extend(_lapiscv_entry(entry, kind) for entry in entries)
    certificates = [str(item).strip() for item in resume.get("certificates") or [] if str(item).strip()]
    if include_certificates and certificates:
        body.append(_lapiscv_section("证书与荣誉"))
        body.append("<p>" + "  ·  ".join(_html_escape(item) for item in certificates) + "</p>")
    return "".join(body)


def _catalog_sidebar(resume: dict, avatar_src: str, *, dark: bool = False) -> str:
    contact_markup = _lapiscv_contact(resume.get("contact") or {})
    overview = _experience_capability_overview(resume)
    avatar_markup = f'<img class="avatar" alt="avatar" src="{avatar_src}">' if avatar_src else ""
    parts = [avatar_markup, f"<h1>{_html_escape(resume.get('name') or '个人简历')}</h1>"]
    if resume.get("title"):
        parts.append(f'<p class="lapiscv-title"><strong>{_html_escape(resume["title"])}</strong></p>')
    if contact_markup:
        parts.append(f'<p class="lapiscv-contact">{contact_markup}</p>')
    if overview:
        parts.extend([_lapiscv_section("经验与能力"), f'<p class="lapiscv-section-summary">{_html_paragraph(overview)}</p>'])
    certificates = [str(item).strip() for item in resume.get("certificates") or [] if str(item).strip()]
    if certificates:
        parts.extend([_lapiscv_section("证书与荣誉"), "<p>" + "<br>".join(_html_escape(item) for item in certificates) + "</p>"])
    suffix = " catalog-sidebar--dark" if dark else ""
    return f'<aside class="catalog-sidebar{suffix}">{"".join(parts)}</aside>'


def resume_to_lapiscv_html(resume: dict, avatar_data: bytes | None = None, design: dict | None = None) -> str:
    """Render an AI resume through the chosen licensed-template layout family."""
    resume = normalize_resume_content(resume)
    avatar_src = _lapiscv_avatar_src(avatar_data)
    variant = str((design or {}).get("layout_variant") or "top_profile")
    if variant not in {"top_profile", "left_sidebar", "right_sidebar", "banner_timeline", "split_columns", "executive_minimal", "creative_asymmetry", "campus_compact"}:
        variant = "top_profile"
    header = _catalog_header(resume, avatar_src, banner=variant in {"banner_timeline", "creative_asymmetry"})
    sections = _catalog_sections(resume)
    if variant == "left_sidebar":
        body = f'<main class="catalog-resume catalog-two-column">{_catalog_sidebar(resume, avatar_src, dark=True)}<section class="catalog-main">{_catalog_sections(resume, include_overview=False, include_certificates=False)}</section></main>'
    elif variant == "right_sidebar":
        body = f'<main class="catalog-resume catalog-two-column catalog-two-column--right"><section class="catalog-main">{header}{_catalog_sections(resume, include_overview=True, include_certificates=False)}</section>{_catalog_sidebar(resume, avatar_src)}</main>'
    elif variant == "banner_timeline":
        body = f'<main class="catalog-resume catalog-resume--timeline">{header}{sections}</main>'
    elif variant == "split_columns":
        overview = _experience_capability_overview(resume)
        body = f'<main class="catalog-resume catalog-resume--split">{header}<div class="catalog-summary-grid"><section>{_lapiscv_section("经验与能力概述") if overview else ""}{f"<p>{_html_paragraph(overview)}</p>" if overview else ""}</section><section>{_catalog_sections({**resume, "experience": [], "projects": []}, include_overview=False)}</section></div>{_catalog_sections({**resume, "education": [], "certificates": []}, include_overview=False)}</main>'
    elif variant == "executive_minimal":
        body = f'<main class="catalog-resume catalog-resume--executive">{header}{sections}</main>'
    elif variant == "creative_asymmetry":
        body = f'<main class="catalog-resume catalog-resume--creative">{header}{sections}</main>'
    elif variant == "campus_compact":
        body = f'<main class="catalog-resume catalog-resume--campus">{header}{sections}</main>'
    else:
        body = f'<main class="catalog-resume catalog-resume--top">{header}{sections}</main>'
    return (
        '<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">'
        f"<title>{_html_escape(resume.get('name') or '个人简历')}的简历</title>"
        f"<style>{_lapiscv_css(design)}</style></head><body>{body}</body></html>"
    )


async def resume_to_lapiscv_pdf(resume: dict, avatar_data: bytes | None = None, design: dict | None = None) -> bytes:
    try:
        from playwright.async_api import async_playwright
    except ImportError as exc:
        raise RuntimeError("LapisCV PDF 渲染需要安装 Playwright") from exc

    html = resume_to_lapiscv_html(resume, avatar_data=avatar_data, design=design)
    async with async_playwright() as playwright:
        browser = await playwright.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
        try:
            page = await browser.new_page(viewport={"width": 794, "height": 1123}, device_scale_factor=1)
            await page.emulate_media(media="print")
            await page.set_content(html, wait_until="domcontentloaded", timeout=20_000)
            return await page.pdf(
                format="A4",
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "13mm", "right": "15mm", "bottom": "13mm", "left": "15mm"},
            )
        finally:
            await browser.close()


def _pdf_styles(font_name: str, theme: dict) -> dict[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "ResumeName",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=23,
            leading=28,
            textColor=colors.HexColor(theme["primary"]),
            alignment=TA_LEFT,
            spaceAfter=2 * mm,
        ),
        "title": ParagraphStyle(
            "ResumeTitle",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor(theme["ink"]),
            spaceAfter=2 * mm,
        ),
        "contact": ParagraphStyle(
            "ResumeContact",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=8.5,
            leading=13,
            textColor=colors.HexColor(theme["muted"]),
        ),
        "section": ParagraphStyle(
            "ResumeSection",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=11.5,
            leading=15,
            textColor=colors.HexColor("#F8FAFC") if theme["colored"] else colors.HexColor(theme["primary"]),
            spaceAfter=0,
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9.2,
            leading=14.5,
            textColor=colors.HexColor(theme["ink"]),
            spaceAfter=1.2 * mm,
        ),
        "entry": ParagraphStyle(
            "ResumeEntry",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.6,
            leading=15,
            textColor=colors.HexColor(theme["ink"]),
        ),
        "date": ParagraphStyle(
            "ResumeDate",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=8.3,
            leading=13,
            alignment=TA_RIGHT,
            textColor=colors.HexColor(theme["muted"]),
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9,
            leading=14,
            leftIndent=6.5 * mm,
            firstLineIndent=-2.5 * mm,
            textColor=colors.HexColor(theme["ink"]),
            spaceAfter=0.8 * mm,
        ),
    }


def _pdf_section(story: list, title: str, styles: dict[str, ParagraphStyle], theme: dict) -> None:
    story.append(Spacer(1, 3 * mm))
    if not theme["colored"]:
        story.append(Paragraph(title, styles["section"]))
        story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#9CA6B4"), spaceAfter=2 * mm))
        return
    heading = Table([[Paragraph(title, styles["section"])]], colWidths=[178 * mm])
    heading.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(theme["primary"])),
                ("LEFTPADDING", (0, 0), (-1, -1), 4 * mm),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4 * mm),
                ("TOPPADDING", (0, 0), (-1, -1), 2.3 * mm),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.1 * mm),
            ]
        )
    )
    story.extend([heading, Spacer(1, 2 * mm)])


def _pdf_entry(story: list, entry: dict, kind: str, styles: dict[str, ParagraphStyle]) -> None:
    heading = Table(
        [
            [
                Paragraph(f"<b>{_escape(_entry_heading(entry, kind))}</b>", styles["entry"]),
                Paragraph(_escape(entry.get("period") or ""), styles["date"]),
            ]
        ],
        colWidths=[135 * mm, 35 * mm],
    )
    heading.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.2 * mm),
            ]
        )
    )
    top = [heading]
    if kind != "work" and entry.get("role") and str(entry.get("role")) != _entry_primary(entry, kind):
        top.append(Paragraph(_escape(str(entry["role"])), styles["title"]))
    story.append(KeepTogether(top))
    for detail in entry.get("details") or []:
        story.append(Paragraph("• " + _escape(str(detail)), styles["bullet"]))
    story.append(Spacer(1, 1.4 * mm))


def resume_to_pdf(resume: dict, avatar_data: bytes | None = None, design: dict | None = None) -> bytes:
    resume = normalize_resume_content(resume)
    theme_id = str((design or {}).get("theme_id") or "executive_navy")
    theme = RESUME_DESIGN_THEMES.get(theme_id, RESUME_DESIGN_THEMES["executive_navy"])
    font_name = _pdf_font_name()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=13 * mm,
        bottomMargin=14 * mm,
        title=f"{resume.get('name', '')}的简历",
        author=resume.get("name") or "",
    )
    styles = _pdf_styles(font_name, theme)
    header_left = [Paragraph(_escape(resume.get("name") or "个人简历"), styles["name"])]
    if resume.get("title"):
        header_left.append(Paragraph(_escape(str(resume["title"])), styles["title"]))
    contact = _contact_text(resume.get("contact") or {})
    if contact:
        header_left.append(Paragraph(_escape(contact), styles["contact"]))
    header_data = [header_left]
    header_widths = [178 * mm]
    if avatar_data:
        avatar_width, avatar_height = _avatar_layout_size(avatar_data)
        header_data.append(PDFImage(BytesIO(avatar_data), width=avatar_width * mm, height=avatar_height * mm))
        header_widths = [142 * mm, 36 * mm]
    header = Table([header_data], colWidths=header_widths)
    header.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4 * mm if theme["colored"] else 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4 * mm if theme["colored"] else 0),
                ("TOPPADDING", (0, 0), (-1, -1), 3 * mm if theme["colored"] else 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3 * mm if theme["colored"] else 1.5 * mm),
                ("ALIGN", (-1, 0), (-1, 0), "RIGHT"),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(theme["soft"]))
                if theme["colored"]
                else ("LINEBELOW", (0, 0), (-1, -1), 0, colors.white),
            ]
        )
    )
    story = [header, Spacer(1, 2 * mm)]
    overview = _experience_capability_overview(resume)
    if overview:
        _pdf_section(story, "经验与能力概述", styles, theme)
        story.append(Paragraph(_escape(overview), styles["body"]))
    for title, key, kind in (
        ("工作经历", "experience", "work"),
        ("项目经历", "projects", "project"),
        ("教育经历", "education", "education"),
    ):
        if resume.get(key):
            _pdf_section(story, title, styles, theme)
            for entry in resume[key]:
                _pdf_entry(story, entry, kind, styles)
    if resume.get("certificates"):
        _pdf_section(story, "证书与荣誉", styles, theme)
        story.append(Paragraph(_escape("  ·  ".join(resume["certificates"])), styles["body"]))

    def add_page_number(canvas, _):
        canvas.saveState()
        if theme["colored"]:
            canvas.setFillColor(colors.HexColor(theme["ribbon"]))
            canvas.rect(0, A4[1] - 5 * mm, A4[0], 5 * mm, stroke=0, fill=1)
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor(theme["muted"]))
        canvas.drawRightString(A4[0] - 16 * mm, 7 * mm, str(canvas.getPageNumber()))
        canvas.restoreState()

    document.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return output.getvalue()


def _escape(value: str) -> str:
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
