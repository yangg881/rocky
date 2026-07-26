import io
import os
import uuid
import zipfile
from pathlib import Path

os.environ.update(
    {
        "STORAGE_BACKEND": "memory",
        "AI_MOCK": "true",
        "JWT_SECRET": "test-secret-that-is-long-enough",
        "ADMIN_USERNAME": "test-admin",
        "ADMIN_PASSWORD": "test-admin-password",
        "SMS_MOCK": "true",
    }
)

from docx import Document  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from PIL import Image  # noqa: E402
from pypdf import PdfReader  # noqa: E402

import app.main as main_module  # noqa: E402
from app.main import app  # noqa: E402

PNG_IMAGE = b"\x89PNG\r\n\x1a\nreview-image"
JPEG_IMAGE = b"\xff\xd8\xff\xe0review-image"
WEBP_IMAGE = b"RIFF\x10\x00\x00\x00WEBPreview-image"


def auth_headers(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def register(client: TestClient, username: str) -> dict:
    phone = f"139{uuid.uuid4().int % 10**8:08d}"
    response = client.post(
        "/resume-ai/api/auth/register",
        json={
            "username": username,
            "phone": phone,
            "code": "123456",
            "password": "secure-pass-123",
            "confirm_password": "secure-pass-123",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()


def jpeg_avatar() -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (240, 300), (72, 103, 150)).save(output, format="JPEG")
    return output.getvalue()


def square_jpeg_avatar() -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (300, 300), (72, 103, 150)).save(output, format="JPEG")
    return output.getvalue()


def test_api_security_headers_and_error_contract() -> None:
    """Public API errors stay machine-readable without losing legacy detail text."""
    with TestClient(app) as client:
        health = client.get("/resume-ai/api/health")
        assert health.status_code == 200
        assert health.headers.get("x-request-id")
        assert "default-src 'self'" in (health.headers.get("content-security-policy") or "")

        invalid = client.post("/resume-ai/api/auth/register", json={"username": "x"})
        assert invalid.status_code == 422
        payload = invalid.json()
        assert payload["code"] == "validation_failed"
        assert payload["request_id"]
        assert payload["detail"] == payload["message"]


def test_web_session_cookie_and_global_session_revocation() -> None:
    """Browsers no longer need a script-readable token, while Android Bearer tokens stay supported."""
    with TestClient(app) as client:
        session = register(client, f"cookie-user-{uuid.uuid4().hex[:8]}")
        assert client.cookies.get("zhiday_session")
        assert client.get("/resume-ai/api/auth/me").status_code == 200

        original_headers = auth_headers(session["token"])
        revoked = client.post("/resume-ai/api/auth/logout-all", headers=original_headers)
        assert revoked.status_code == 204
        assert client.get("/resume-ai/api/auth/me", headers=original_headers).status_code == 401


def docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(values)


def test_full_user_and_admin_flow() -> None:
    with TestClient(app) as client:
        suffix = uuid.uuid4().hex[:8]
        first = register(client, f"求职者{suffix}")
        second = register(client, f"另一用户{suffix}")
        first_headers = auth_headers(first["token"])
        second_headers = auth_headers(second["token"])

        mismatch = client.post(
            "/resume-ai/api/auth/register",
            json={
                "username": f"失败{suffix}",
                "phone": f"138{uuid.uuid4().int % 10**8:08d}",
                "code": "123456",
                "password": "secure-pass-123",
                "confirm_password": "different-pass",
            },
        )
        assert mismatch.status_code == 422

        resume_payload = {
            "name": "产品经理主简历",
            "content": {
                "name": "张三",
                "title": "产品经理",
                "contact": {"email": "test@example.com"},
                "summary": "负责企业产品规划与交付",
                "skills": ["需求分析", "项目管理"],
                "experience": [
                    {"company": "早期公司", "role": "专员", "period": "2018.01-2020.02", "details": "负责早期项目"},
                    {
                        "company": "示例公司",
                        "role": "产品经理",
                        "period": "2024.05-2026.01",
                        "details": "负责真实项目交付",
                    },
                    {"company": "中期公司", "role": "工程师", "period": "2022.12-2023.07", "details": "负责中期项目"},
                    {"company": "当前公司", "role": "负责人", "period": "2026.02-至今", "details": "负责当前项目"},
                ],
                "projects": [],
                "education": [],
                "certificates": [],
            },
        }
        created = client.post("/resume-ai/api/resumes", json=resume_payload, headers=first_headers)
        assert created.status_code == 201, created.text
        resume = created.json()
        assert resume["is_default"] is True

        avatar = client.post(
            f"/resume-ai/api/resumes/{resume['id']}/avatar",
            files={"file": ("avatar.jpg", jpeg_avatar(), "image/jpeg")},
            headers=first_headers,
        )
        assert avatar.status_code == 200, avatar.text
        assert avatar.json()["avatar_key"].endswith(f"/{resume['id']}.jpg")

        isolated = client.get(f"/resume-ai/api/resumes/{resume['id']}", headers=second_headers)
        assert isolated.status_code == 404

        jd_response = client.post(
            "/resume-ai/api/jd/parse",
            json={
                "source_type": "text",
                "text": "招聘高级产品经理，负责需求分析、项目管理和跨团队协作，要求三年以上相关经验。",
            },
            headers=first_headers,
        )
        assert jd_response.status_code == 202, jd_response.text
        assert jd_response.json()["status"] == "processing"
        jd_task = client.get(
            f"/resume-ai/api/jd/tasks/{jd_response.json()['id']}", headers=first_headers
        )
        assert jd_task.status_code == 200
        assert jd_task.json()["status"] == "completed"
        jd = jd_task.json()["result"]

        failed_task_id = str(uuid.uuid4())
        main_module.store.put_json(
            "a",
            f"jd-tasks/{first['user']['id']}/{failed_task_id}.json",
            {
                "id": failed_task_id,
                "user_id": first["user"]["id"],
                "source": "url",
                "source_detail": "https://example.com/failed-job",
                "status": "failed",
                "result": None,
                "error": "测试失败原因",
                "created_at": "9999-01-01T00:00:00+00:00",
                "updated_at": "9999-01-01T00:00:00+00:00",
            },
        )
        jd_history = client.get("/resume-ai/api/jd/tasks", headers=first_headers)
        assert jd_history.status_code == 200
        assert jd_history.json()[0]["id"] == failed_task_id
        assert jd_history.json()[0]["error"] == "测试失败原因"
        assert any(item.get("result") == jd for item in jd_history.json())
        assert client.get("/resume-ai/api/jd/tasks", headers=second_headers).json() == []

        multi_image_jd = client.post(
            "/resume-ai/api/jd/ocr",
            files=[
                ("files", ("page-1.png", PNG_IMAGE + b"-first", "image/png")),
                ("files", ("page-2.jpeg", JPEG_IMAGE, "image/jpeg")),
                ("files", ("page-3.webp", WEBP_IMAGE, "image/webp")),
            ],
            headers=first_headers,
        )
        assert multi_image_jd.status_code == 202, multi_image_jd.text
        image_task = client.get(
            f"/resume-ai/api/jd/tasks/{multi_image_jd.json()['id']}", headers=first_headers
        ).json()
        assert image_task["status"] == "completed"
        assert image_task["result"]["image_count"] == 3
        assert image_task["result"]["source_keys"][0].endswith("-01.png")
        assert image_task["result"]["source_keys"][1].endswith("-02.jpg")
        assert image_task["result"]["source_keys"][2].endswith("-03.webp")

        too_many_images = client.post(
            "/resume-ai/api/jd/ocr",
            files=[("files", (f"page-{index}.png", b"image", "image/png")) for index in range(11)],
            headers=first_headers,
        )
        assert too_many_images.status_code == 413

        generation_response = client.post(
            "/resume-ai/api/generations",
            json={"resume_id": resume["id"], "jd": jd},
            headers=first_headers,
        )
        assert generation_response.status_code == 202, generation_response.text
        queued = generation_response.json()
        assert queued["status"] == "processing"
        assert queued["files"] == {}
        assert queued["design"]["theme_id"] == "tech_indigo"
        assert "cover" not in queued

        history = client.get("/resume-ai/api/generations", headers=first_headers)
        assert history.status_code == 200
        assert len(history.json()) == 1
        generation = history.json()[0]
        assert generation["status"] == "completed"
        assert generation["files"]["docx"]["size"] > 0
        assert generation["files"]["pdf"]["size"] > 0
        assert generation["design"]["theme_id"] == "tech_indigo"
        assert generation["design"]["source"] == "ai"
        assert "cover" not in generation

        docx_download = client.get(
            f"/resume-ai/api/generations/{generation['id']}/download/docx", headers=first_headers
        )
        assert docx_download.status_code == 200
        assert docx_download.content.startswith(b"PK")
        assert "attachment" in docx_download.headers["content-disposition"]
        document_text = docx_text(docx_download.content)
        assert "张三" in document_text
        assert "工作经历" in document_text
        assert "经验与能力概述" in document_text
        assert "个人概述" not in document_text
        assert "经验与能力\n" not in document_text
        assert "专业技能" not in document_text
        assert "示例公司  产品经理" in document_text
        document_company_positions = [
            document_text.index(company) for company in ("当前公司", "示例公司", "中期公司", "早期公司")
        ]
        assert document_company_positions == sorted(document_company_positions)
        with zipfile.ZipFile(io.BytesIO(docx_download.content)) as archive:
            assert any(name.startswith("word/media/") for name in archive.namelist())

        pdf_download = client.get(
            f"/resume-ai/api/generations/{generation['id']}/download/pdf", headers=first_headers
        )
        assert pdf_download.status_code == 200
        assert pdf_download.headers["content-type"].startswith("application/pdf")
        assert pdf_download.content.startswith(b"%PDF")
        pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(pdf_download.content)).pages)
        assert "张三" in pdf_text
        assert "工作经历" in pdf_text
        assert "经验与能力概述" in pdf_text
        assert "个人概述" not in pdf_text
        assert "示例公司 产品经理" in pdf_text
        pdf_company_positions = [
            pdf_text.index(company) for company in ("当前公司", "示例公司", "中期公司", "早期公司")
        ]
        assert pdf_company_positions == sorted(pdf_company_positions)

        print_friendly = client.post(
            "/resume-ai/api/generations",
            json={"resume_id": resume["id"], "jd": jd, "design_theme": "ats_mono"},
            headers=first_headers,
        )
        assert print_friendly.status_code == 202, print_friendly.text
        print_generation = next(
            item
            for item in client.get("/resume-ai/api/generations", headers=first_headers).json()
            if item["id"] == print_friendly.json()["id"]
        )
        assert print_generation["status"] == "completed"
        assert print_generation["design"]["theme_id"] == "ats_mono"
        assert print_generation["design"]["source"] == "manual"

        themed_again = client.post(
            f"/resume-ai/api/generations/{generation['id']}/regenerate",
            json={"design_theme": "creative_plum"},
            headers=first_headers,
        )
        assert themed_again.status_code == 202, themed_again.text
        regenerated = next(
            item
            for item in client.get("/resume-ai/api/generations", headers=first_headers).json()
            if item["id"] == themed_again.json()["id"]
        )
        assert regenerated["status"] == "completed"
        assert regenerated["parent_generation_id"] == generation["id"]
        assert regenerated["design"]["theme_id"] == "creative_plum"
        assert regenerated["optimized"] == generation["optimized"]

        file_link = client.get(
            "/resume-ai/api/file-link",
            params={"bucket": "b", "key": generation["files"]["pdf"]["key"]},
            headers=first_headers,
        )
        assert file_link.status_code == 200
        assert file_link.json()["url"].startswith("memory://")

        admin_login = client.post(
            "/resume-ai/api/auth/login", json={"username": "test-admin", "password": "test-admin-password"}
        )
        assert admin_login.status_code == 200
        admin_headers = auth_headers(admin_login.json()["token"])
        stats = client.get("/resume-ai/api/admin/stats", headers=admin_headers)
        assert stats.status_code == 200
        assert stats.json()["users"] >= 2
        assert stats.json()["generations"] >= 1

        admin_docx = client.get(
            f"/resume-ai/api/admin/generations/{generation['user_id']}/{generation['id']}/download/docx",
            headers=admin_headers,
        )
        assert admin_docx.status_code == 200
        assert admin_docx.content.startswith(b"PK")
        admin_pdf = client.get(
            f"/resume-ai/api/admin/generations/{generation['user_id']}/{generation['id']}/download/pdf",
            headers=admin_headers,
        )
        assert admin_pdf.status_code == 200
        assert admin_pdf.content.startswith(b"%PDF")

        tasks = client.get("/resume-ai/api/admin/tasks", headers=admin_headers)
        assert tasks.status_code == 200
        task_types = {item["task_type"] for item in tasks.json()}
        assert {"jd_parse", "resume_generation"}.issubset(task_types)

        normal_user_denied = client.get("/resume-ai/api/admin/users", headers=first_headers)
        assert normal_user_denied.status_code == 403

        new_phone = f"137{uuid.uuid4().int % 10**8:08d}"
        admin_phone = client.patch(
            f"/resume-ai/api/admin/users/{second['user']['id']}/phone",
            json={"phone": new_phone},
            headers=admin_headers,
        )
        assert admin_phone.status_code == 200
        assert admin_phone.json()["user"]["phone"] == new_phone

        deleted = client.delete(f"/resume-ai/api/admin/users/{first['user']['id']}", headers=admin_headers)
        assert deleted.status_code == 200
        remaining_tasks = client.get("/resume-ai/api/admin/tasks", headers=admin_headers).json()
        assert all(item["user_id"] != first["user"]["id"] for item in remaining_tasks)


def test_nested_model_result_still_exports_content(monkeypatch) -> None:
    async def wrapped_resume(_resume: dict, _jd: dict) -> dict:
        return {
            "resume": {
                "name": "李四",
                "title": "高级产品经理",
                "summary": "推动复杂产品从规划到交付",
                "skills": ["需求分析", "项目管理"],
                "experience": [{"company": "真实公司", "role": "产品经理", "details": ["负责核心产品交付"]}],
            }
        }

    monkeypatch.setattr(main_module.ai, "rewrite_resume", wrapped_resume)
    with TestClient(app) as client:
        user = register(client, f"嵌套结果{uuid.uuid4().hex[:8]}")
        headers = auth_headers(user["token"])
        created = client.post(
            "/resume-ai/api/resumes",
            json={"name": "基础简历", "content": {"name": "李四", "experience": []}},
            headers=headers,
        )
        response = client.post(
            "/resume-ai/api/generations",
            json={"resume_id": created.json()["id"], "jd": {"title": "高级产品经理"}},
            headers=headers,
        )
        assert response.status_code == 202
        generation = client.get("/resume-ai/api/generations", headers=headers).json()[0]
        assert generation["status"] == "completed"
        download = client.get(
            f"/resume-ai/api/generations/{generation['id']}/download/docx", headers=headers
        )
        document_text = docx_text(download.content)
        assert "李四" in document_text
        assert "真实公司" in document_text
        assert "工作经历" in document_text


def test_generation_failure_is_recorded(monkeypatch) -> None:
    async def fail_rewrite(_resume: dict, _jd: dict) -> dict:
        raise RuntimeError("simulated model failure")

    monkeypatch.setattr(main_module.ai, "rewrite_resume", fail_rewrite)
    with TestClient(app) as client:
        user = register(client, f"失败任务{uuid.uuid4().hex[:8]}")
        headers = auth_headers(user["token"])
        created = client.post(
            "/resume-ai/api/resumes",
            json={"name": "基础简历", "content": {"name": "王五"}},
            headers=headers,
        )
        response = client.post(
            "/resume-ai/api/generations",
            json={"resume_id": created.json()["id"], "jd": {"title": "运营经理"}},
            headers=headers,
        )
        assert response.status_code == 202
        generation = client.get("/resume-ai/api/generations", headers=headers).json()[0]
        assert generation["status"] == "failed"
        assert generation["files"] == {}
        blocked_download = client.get(
            f"/resume-ai/api/generations/{generation['id']}/download/pdf", headers=headers
        )
        assert blocked_download.status_code == 409


def test_jd_failure_is_recorded_with_reason(monkeypatch) -> None:
    async def fail_jd(*_args, **_kwargs) -> dict:
        raise main_module.AIServiceError(
            "模型暂时不可用",
            category="output_truncated",
            metadata={"model": "test-model", "finish_reason": "length", "retry_count": 1},
        )

    monkeypatch.setattr(main_module.ai, "structure_jd", fail_jd)
    with TestClient(app) as client:
        user = register(client, f"岗位失败{uuid.uuid4().hex[:8]}")
        headers = auth_headers(user["token"])
        queued = client.post(
            "/resume-ai/api/jd/parse",
            json={"source_type": "text", "text": "招聘运营经理，负责内容运营和活动策划，要求相关工作经验。"},
            headers=headers,
        )
        assert queued.status_code == 202
        task = client.get(f"/resume-ai/api/jd/tasks/{queued.json()['id']}", headers=headers).json()
        assert task["status"] == "failed"
        assert task["error"] == "模型暂时不可用"

        admin_login = client.post(
            "/resume-ai/api/auth/login", json={"username": "test-admin", "password": "test-admin-password"}
        )
        admin_headers = auth_headers(admin_login.json()["token"])
        admin_tasks = client.get("/resume-ai/api/admin/tasks", headers=admin_headers).json()
        failed = next(item for item in admin_tasks if item["id"] == queued.json()["id"])
        assert failed["status"] == "failed"
        assert failed["error"] == "模型暂时不可用"
        assert failed["error_category"] == "output_truncated"
        assert failed["model_metadata"]["finish_reason"] == "length"


def test_empty_jd_parse_result_is_failed_and_cannot_generate(monkeypatch) -> None:
    async def empty_jd(*_args, **_kwargs) -> dict:
        return {
            "title": "",
            "company": "",
            "responsibilities": [],
            "requirements": [],
            "preferred": [],
            "keywords": [],
            "raw_text": "只有网页导航和广告，没有岗位正文",
        }

    monkeypatch.setattr(main_module.ai, "structure_jd", empty_jd)
    with TestClient(app) as client:
        user = register(client, f"空岗位{uuid.uuid4().hex[:8]}")
        headers = auth_headers(user["token"])
        queued = client.post(
            "/resume-ai/api/jd/parse",
            json={
                "source_type": "text",
                "text": "这是一段长度足够但没有任何岗位职位职责要求的无效内容，只包含页面提示和广告。",
            },
            headers=headers,
        )
        assert queued.status_code == 202
        task = client.get(f"/resume-ai/api/jd/tasks/{queued.json()['id']}", headers=headers).json()
        assert task["status"] == "failed"
        assert "岗位解析结果为空" in task["error"]

        created = client.post(
            "/resume-ai/api/resumes",
            json={"name": "基础简历", "content": {"name": "测试用户"}},
            headers=headers,
        )
        blocked = client.post(
            "/resume-ai/api/generations",
            json={
                "resume_id": created.json()["id"],
                "jd": {
                    "title": "",
                    "company": "",
                    "responsibilities": [],
                    "requirements": [],
                    "preferred": [],
                    "keywords": [],
                    "raw_text": "只有网页导航和广告，没有岗位正文",
                },
            },
            headers=headers,
        )
        assert blocked.status_code == 422


def test_jd_image_validation_legacy_field_and_partial_write_cleanup(monkeypatch) -> None:
    with TestClient(app, raise_server_exceptions=False) as client:
        user = register(client, f"截图用户{uuid.uuid4().hex[:8]}")
        headers = auth_headers(user["token"])

        unsupported = client.post(
            "/resume-ai/api/jd/ocr",
            files={"file": ("job.svg", b"<svg/>", "image/svg+xml")},
            headers=headers,
        )
        assert unsupported.status_code == 400

        mismatched_magic = client.post(
            "/resume-ai/api/jd/ocr",
            files={"file": ("job.png", b"not-a-png", "image/png")},
            headers=headers,
        )
        assert mismatched_magic.status_code == 400

        legacy = client.post(
            "/resume-ai/api/jd/ocr",
            files={"file": ("legacy.png", PNG_IMAGE, "image/png")},
            headers=headers,
        )
        assert legacy.status_code == 202, legacy.text
        legacy_task = client.get(f"/resume-ai/api/jd/tasks/{legacy.json()['id']}", headers=headers).json()
        assert legacy_task["status"] == "completed"
        assert legacy_task["result"]["image_count"] == 1
        assert legacy_task["result"]["source_key"] == legacy_task["result"]["source_keys"][0]

        def jd_image_keys() -> set[tuple[str, str]]:
            return {item for item in main_module.store._memory if "/jd-images/" in item[1]}

        keys_before = jd_image_keys()
        original_put_bytes = main_module.store.put_bytes
        jd_write_count = 0

        def fail_second_jd_write(bucket: str, key: str, data: bytes, content_type: str) -> None:
            nonlocal jd_write_count
            if "/jd-images/" in key:
                jd_write_count += 1
                if jd_write_count == 2:
                    raise RuntimeError("simulated storage failure")
            original_put_bytes(bucket, key, data, content_type)

        monkeypatch.setattr(main_module.store, "put_bytes", fail_second_jd_write)
        partial_failure = client.post(
            "/resume-ai/api/jd/ocr",
            files=[
                ("files", ("page-1.png", PNG_IMAGE + b"-one", "image/png")),
                ("files", ("page-2.png", PNG_IMAGE + b"-two", "image/png")),
            ],
            headers=headers,
        )
        assert partial_failure.status_code == 500
        assert jd_image_keys() == keys_before


def test_self_service_password_reset() -> None:
    with TestClient(app) as client:
        username = f"重置用户{uuid.uuid4().hex[:8]}"
        user = register(client, username)
        phone = user["user"]["phone"]
        sent = client.post("/resume-ai/api/auth/sms-code", json={"phone": phone, "scene": "reset_password"})
        assert sent.status_code == 200
        reset = client.post(
            "/resume-ai/api/auth/reset-password",
            json={
                "phone": phone,
                "code": "123456",
                "new_password": "new-secure-pass",
                "confirm_password": "new-secure-pass",
            },
        )
        assert reset.status_code == 200
        login = client.post(
            "/resume-ai/api/auth/login", json={"username": username, "password": "new-secure-pass"}
        )
        assert login.status_code == 200


def test_sms_login_only_allows_registered_phone() -> None:
    with TestClient(app) as client:
        registered = register(client, f"验证码登录{uuid.uuid4().hex[:8]}")
        phone = registered["user"]["phone"]
        sent = client.post("/resume-ai/api/auth/sms-code", json={"phone": phone, "scene": "login"})
        assert sent.status_code == 200
        login = client.post("/resume-ai/api/auth/sms-login", json={"phone": phone, "code": "123456"})
        assert login.status_code == 200
        assert login.json()["user"]["id"] == registered["user"]["id"]

        unknown = client.post(
            "/resume-ai/api/auth/sms-code",
            json={"phone": f"137{uuid.uuid4().int % 10**8:08d}", "scene": "login"},
        )
        assert unknown.status_code == 404


def test_sms_phone_registration_and_change_phone() -> None:
    with TestClient(app) as client:
        phone = f"136{uuid.uuid4().int % 10**8:08d}"
        sent = client.post("/resume-ai/api/auth/sms-code", json={"phone": phone, "scene": "register"})
        assert sent.status_code == 200
        response = client.post(
            "/resume-ai/api/auth/register",
            json={
                "username": f"短信用户{uuid.uuid4().hex[:8]}",
                "phone": phone,
                "code": "123456",
                "password": "secure-pass-123",
                "confirm_password": "secure-pass-123",
            },
        )
        assert response.status_code == 201
        headers = auth_headers(response.json()["token"])
        assert response.json()["user"]["phone"] == phone

        duplicate = client.post(
            "/resume-ai/api/auth/register",
            json={
                "username": f"重复手机{uuid.uuid4().hex[:8]}",
                "phone": phone,
                "code": "123456",
                "password": "secure-pass-123",
                "confirm_password": "secure-pass-123",
            },
        )
        assert duplicate.status_code == 409

        new_phone = f"135{uuid.uuid4().int % 10**8:08d}"
        sent_change = client.post(
            "/resume-ai/api/auth/sms-code",
            json={"phone": new_phone, "scene": "change_phone"},
            headers=headers,
        )
        assert sent_change.status_code == 200
        changed = client.post(
            "/resume-ai/api/auth/change-phone",
            json={"phone": new_phone, "code": "123456"},
            headers=headers,
        )
        assert changed.status_code == 200
        assert changed.json()["user"]["phone"] == new_phone


def test_self_service_account_deletion_removes_user_data_and_indexes() -> None:
    with TestClient(app) as client:
        username = f"注销用户{uuid.uuid4().hex[:8]}"
        registered = register(client, username)
        headers = auth_headers(registered["token"])
        user_id = registered["user"]["id"]
        phone = registered["user"]["phone"]
        created = client.post(
            "/resume-ai/api/resumes",
            json={"name": "待删除简历", "content": {"name": "测试用户"}},
            headers=headers,
        )
        assert created.status_code == 201
        avatar = client.post(
            f"/resume-ai/api/resumes/{created.json()['id']}/avatar",
            files={"file": ("avatar.jpg", square_jpeg_avatar(), "image/jpeg")},
            headers=headers,
        )
        assert avatar.status_code == 200
        deleted = client.post(
            "/resume-ai/api/auth/delete-account",
            json={"current_password": "secure-pass-123", "confirm_username": username},
            headers=headers,
        )
        assert deleted.status_code == 200
        assert client.get("/resume-ai/api/auth/me", headers=headers).status_code == 401
        assert main_module.store.get_json("a", f"users/{user_id}.json") is None
        assert main_module.get_user_by_username(username) is None
        assert main_module.get_user_by_phone(phone) is None
        assert not main_module.store.list_json("a", f"resumes/{user_id}/")
        assert not any(user_id in key for _, key in main_module.store._memory)


def test_square_avatar_keeps_aspect_ratio_in_docx_and_pdf() -> None:
    with TestClient(app) as client:
        user = register(client, f"方形头像{uuid.uuid4().hex[:8]}")
        headers = auth_headers(user["token"])
        resume = client.post(
            "/resume-ai/api/resumes",
            json={"name": "方形头像简历", "content": {"name": "头像测试", "summary": "真实经历"}},
            headers=headers,
        ).json()
        client.post(
            f"/resume-ai/api/resumes/{resume['id']}/avatar",
            files={"file": ("square.jpg", square_jpeg_avatar(), "image/jpeg")},
            headers=headers,
        )
        queued = client.post(
            "/resume-ai/api/generations",
            json={"resume_id": resume["id"], "jd": {"title": "测试岗位"}},
            headers=headers,
        ).json()
        docx_data = client.get(
            f"/resume-ai/api/generations/{queued['id']}/download/docx", headers=headers
        ).content
        document = Document(io.BytesIO(docx_data))
        assert len(document.inline_shapes) == 1
        shape = document.inline_shapes[0]
        assert abs(shape.width / shape.height - 1) < 0.01
        assert abs(shape.height / 36000 - 30) < 0.1
        pdf_data = client.get(
            f"/resume-ai/api/generations/{queued['id']}/download/pdf", headers=headers
        ).content
        first_image = PdfReader(io.BytesIO(pdf_data)).pages[0].images[0]
        with Image.open(io.BytesIO(first_image.data)) as rendered_avatar:
            assert abs(rendered_avatar.width / rendered_avatar.height - 1) < 0.01


def test_document_preview_assets_are_served_locally() -> None:
    with TestClient(app) as client:
        pdf_viewer = client.get("/resume-ai/static/vendor/pdf.min.mjs")
        pdf_worker = client.get("/resume-ai/static/vendor/pdf.worker.min.mjs")
        pdf_worker_compat = client.get("/resume-ai/static/vendor/pdf.worker.compat.mjs")
        docx_viewer = client.get("/resume-ai/static/vendor/docx-preview.min.js")
        jszip = client.get("/resume-ai/static/vendor/jszip.min.js")
        assert pdf_viewer.status_code == 200
        assert pdf_worker.status_code == 200
        assert pdf_worker_compat.status_code == 200
        assert docx_viewer.status_code == 200
        assert jszip.status_code == 200
        assert len(pdf_viewer.content) > 300_000
        assert len(pdf_worker.content) > 1_000_000
        assert b"Promise.withResolvers" in pdf_worker_compat.content
        assert b"renderAsync" in docx_viewer.content
        assert b"JSZip" in jszip.content


def test_document_preview_frontend_supports_admin_and_mobile_pdf() -> None:
    user_script = (main_module.STATIC_DIR / "app.js").read_text(encoding="utf-8")
    admin_script = (main_module.STATIC_DIR / "admin.js").read_text(encoding="utf-8")
    admin_html = (main_module.STATIC_DIR / "admin.html").read_text(encoding="utf-8")
    assert "Promise.withResolvers" in user_script
    assert "Promise.withResolvers" in admin_script
    assert "pdf.worker.compat.mjs" in user_script
    assert "pdf.worker.compat.mjs" in admin_script
    assert 'data-record-preview="docx"' in admin_script
    assert 'data-record-preview="pdf"' in admin_script
    assert "docx-preview.min.js" in admin_script
    assert "pdf.min.mjs" in admin_script
    assert "admin-preview-dialog" in admin_html


def test_frontend_blocks_empty_jd_generation_records() -> None:
    user_script = (main_module.STATIC_DIR / "app.js").read_text(encoding="utf-8")
    assert "function isUsableJd" in user_script
    assert "解析无效" in user_script
    assert "这条解析记录没有可用的岗位结果" in user_script


def test_pdf_page_decoration_does_not_use_left_side_stripe() -> None:
    source = (Path(__file__).resolve().parents[1] / "app" / "services.py").read_text(encoding="utf-8")
    assert "A4[1] - 52 * mm" not in source


def test_admin_task_table_markup_keeps_reason_cell_closed() -> None:
    script = (main_module.STATIC_DIR / "admin.js").read_text(encoding="utf-8")
    assert 'class="task-reason ${task.error ? "error" : ""}">' in script
